"""Reports functions."""
import os
import re
import csv
import uuid
import base64
import time
import string
import calendar
import collections
import pandas
import memcache
from django.db import connection
from django.db import connections
from datetime import datetime, timedelta, date
from calendar import monthrange, month_name
from collections import OrderedDict
from openpyxl import Workbook
from reportlab.pdfgen import canvas
from reportlab.lib.enums import TA_JUSTIFY, TA_RIGHT, TA_CENTER, TA_LEFT
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Image, Frame,
    Table, TableStyle)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm, mm
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.lib.pagesizes import A4, letter, landscape
from reportlab.lib.utils import ImageReader
# Security
from reportlab.graphics.barcode import qr
from reportlab.graphics.shapes import Drawing
from reportlab.graphics import renderPDF
import openpyxl
import pandas as pd
import numpy as np
from openpyxl.styles import colors, PatternFill

from cpovc_main.functions import (
    get_general_list, get_dict, get_mapped, convert_date)
from cpovc_main.models import SetupGeography

from cpovc_ovc.models import (
    OVCAggregate, OVCRegistration, OVCCluster, OVCClusterCBO)

from .config import reports
from cpovc_registry.models import (
    RegOrgUnitGeography, RegPerson, RegOrgUnit, RegPersonsSiblings,
    RegPersonsGeo, RegPersonsGuardians, RegPersonsExternalIds,
    RegPersonsOrgUnits)

from cpovc_forms.models import (
    OVCCaseCategory, OVCCaseGeo, OVCCaseEventServices,
    OVCPlacement, OVCAdverseEventsOtherFollowUp, OVCCaseEventClosure,
    OVCDischargeFollowUp, OVCCaseRecord, OVCAdverseEventsFollowUp,
    OVCCaseEvents)
from cpovc_auth.models import AppUser

from django.conf import settings
from django.db.models import Count
from .queries import QUERIES, REPORTS
from .parameters import ORPTS, RPTS, GRPTS, ADHC
from .security import BarCode

from reportlab.lib import colors


MEDIA_ROOT = settings.MEDIA_ROOT
STATIC_ROOT = settings.STATICFILES_DIRS[0]
DOC_ROOT = settings.DOCUMENT_ROOT


class Canvas(canvas.Canvas):
    """Pagination extention for canvas."""

    def __init__(self, *args, **kwargs):
        """Constructor for my pagination."""
        canvas.Canvas.__init__(self, *args, **kwargs)
        self._saved_page_states = []
        self.page = ''

    def showPage(self):
        """Get the pages first."""
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        """To add page info to each page (page x of y)."""
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_number(num_pages)
            canvas.Canvas.showPage(self)
        canvas.Canvas.save(self)

    def draw_page_number(self, page_count):
        """Draw the final page."""
        self.page = "Page %d of %d" % (self._pageNumber, page_count)
        self.setFont("Helvetica", 8)
        self.drawString(0.5 * inch, 0.5 * inch, self.page)
        # self.drawRightString(70, 20, self.page)
        '''
        self.drawRightString(20 * cm, 1 * cm,
                             "Page %d of %d" % (self._pageNumber, page_count))
        '''


def get_report_body(params, report='DSCE'):
    """Method to get report body."""
    try:
        fields = OrderedDict()
        dot_padd = '.' * 170
        dash_padd = ''
        # '_' * 70
        params = child_data(params)
        values = {'padd_dash': dash_padd, 'padd_dot': dot_padd,
                  'child_county': '', 'nationality': 'KENYAN',
                  'religion': '', 'child_sub_county': '', 'source': 'CPIMS',
                  'child_ward': '', 'child_chief': '', 'child_village': ''}
        # params = merge_two_dicts(values, rparams)
        for item in values:
            if item not in params:
                params[item] = values[item]
        report_id = report if report in reports else 'NONE'
        report_params = reports[report_id]
        report_data = report_params % params
        repos = report_data.split('\n')
        cnt = 0
        for repo_val in repos:
            field_values = repo_val.strip()
            field_vals = field_values.split(':', 1)
            field_txt = None if len(field_vals) < 2 else field_vals[1]
            chf = field_vals[0]
            if '<' in chf and '>' in chf and 'br/' not in chf:
                cnt += 1
                fields[field_vals[0].replace('>', '_%s>' % cnt)] = field_txt
            else:
                fields[field_vals[0]] = field_txt
        return fields
    except Exception as e:
        raise e


def get_case_details(field_names):
    """Method to get only case categories from list general."""
    try:
        case_categories = get_general_list(field_names)
    except Exception as e:
        print("Error getting case categories - %s" % (str(e)))
        pass
    else:
        return case_categories


def case_load_header(report_type=1, header=False, params=[]):
    """Method to generate reports headings."""
    try:
        report_region = 0
        age_ranges = ['0 - 5 yrs', '6 - 10 yrs', '11 - 15 yrs',
                      '16 - 17 yrs', '18+ yrs', 'Sub-Total']
        genders = ['M', 'F']
        html = ""
        titles = {1: 'Case Category (New Cases)',
                  2: 'Summary of {period_name}',
                  3: 'List of Diseases', 4: 'Situation of Children'}
        report_type_id = report_type if report_type in titles else 1
        title = titles[report_type_id]
        # Params filter
        if params:
            rrg = 'report_region'
            report_region = params[rrg] if rrg in params else 0
        if header:
            html = ("<tr><td colspan='3'>County</td><td colspan='6'>"
                    "{county}</td><td colspan='3'>Year</td>"
                    "<td colspan='4'>{years}</td></tr>")
            html += ("<tr><td colspan='3'>Sub County</td><td colspan='6'>"
                     "{sub_county}</td><td colspan='3'>Month</td>"
                     "<td colspan='4'>{month}</td></tr>")
            if report_region == 4:
                html = ("<tr><td colspan='3'>Organisation</td><td colspan='6'>"
                        "{org_unit_name}</td><td colspan='3'>Year</td>"
                        "<td colspan='4'>{years}</td></tr>")
                html += ("<tr><td colspan='3'>Unit</td><td colspan='6'>"
                         "{org_units_name}</td><td colspan='3'>Month</td>"
                         "<td colspan='4'>{month}</td></tr>")
            html += ("<tr><td colspan='16'>{period} Case Load Summary"
                     "</td></tr>")
        html += "<tr><td colspan='3' rowspan='2'>%s</td>" % (title)
        for age_range in age_ranges:
            html += "<td colspan='2'>%s</td>" % (age_range)
        html += "<td rowspan='2'>TOTAL</td></tr>"
        html += "<tr>"
        for age_range in age_ranges:
            for gender in genders:
                html += "<td>%s</td>" % (gender)
        html += "</tr>"
        return html
    except Exception as e:
        raise e


def get_data_element(item='case'):
    """Method to get data elements other than categories and inteventions."""
    results = {}
    try:
        case, params = {}, {}
        case[1] = 'Total Children'
        case[2] = 'Total Cases'
        case[3] = 'Total Interventions'
        case[4] = 'Percentage Interventions'
        # dropped out -  (90+ days no intervention)
        case[5] = 'Dropped Out'
        case[6] = 'Pending'

        results['case'] = case
        if item in results:
            params = results[item]
    except Exception:
        pass
    else:
        return params


def simple_documents(params, document_name='CPIMS', report_name='letter'):
    """Method to generate simple report."""
    try:
        file_name = "%s/%s.pdf" % (MEDIA_ROOT, report_name)
        doc = SimpleDocTemplate(file_name, pagesize=letter,
                                rightMargin=48, leftMargin=48,
                                topMargin=30, bottomMargin=30)
        story = []
        logo = "%s/img/gok_logo.jpg" % (STATIC_ROOT)
        today = datetime.now()
        d = int(today.strftime('%d'))
        st_rd = {1: 'st', 2: 'nd', 3: 'rd'}
        suffix = 'th' if 11 <= d <= 13 else st_rd.get(d % 10, 'th')
        # formatted_date = today.strftime('%d{S} %B %Y').replace('{S}', suffix)
        suffix = '<sup>%s</sup>' % (suffix)
        formatted_date = '{dt.day}{S} {dt:%B}, {dt.year}'.format(
            dt=today, S=suffix)

        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name='Justify', alignment=TA_JUSTIFY, leading=22))

        styles.add(ParagraphStyle(name="Centered", alignment=TA_CENTER,
                                  leading=22))
        heading = 'Ministry of Labour and East African Community Affairs'
        sub_heading = 'DEPARTMENT OF CHILDREN SERVICES'
        ptext = '<font size=16><b>%s</b></font>' % heading.upper()
        story.append(Paragraph(ptext, styles["Centered"]))
        # story.append(Spacer(1, 12))
        ptext = '<font size=13><b>%s</b></font>' % sub_heading
        story.append(Paragraph(ptext, styles["Centered"]))
        story.append(Spacer(1, 12))

        sst = getSampleStyleSheet()
        tgram = '.' * 34
        ref = '.' * 36

        p0 = Paragraph('''
            <font size=9>Telegram %s <br/>
            Telephone: 020 2077160<br/>
            When replying please quote<br/><br/>
            Ref. No. %s</font>
            ''' % (tgram, ref), sst["BodyText"])

        org_unit = params['org_unit']
        address = params['address']
        p1 = Paragraph('''
               <font size=9>%s<br/>
               %s<br/><br/>
               Date. %s</font>
               ''' % (org_unit, address, formatted_date), sst["BodyText"])
        im = Image(logo, 1.2 * inch, 1.0 * inch)
        # story.append(im)

        data = [[p0, im, p1]]

        t = Table(data, style=[('ALIGN', (0, 0), (2, 0), 'CENTER'),
                               ('VALIGN', (0, 0), (2, 0), 'MIDDLE'), ])

        story.append(t)
        story.append(Spacer(1, 12))

        # Create return address
        ptext = ('<para align=center spaceb=3><font size=12>'
                 '<b>%s</b></font></para>') % document_name.upper()
        story.append(Paragraph(ptext, styles["Normal"]))
        story.append(Spacer(1, 24))
        vals = get_report_body(params, report_name)
        ptext = '%s' % (vals)
        story.append(Paragraph(ptext, styles["Justify"]))
        story.append(Spacer(1, 12))
        doc.build(story, onFirstPage=draw_page, onLaterPages=draw_page,
                  canvasmaker=Canvas)
    except Exception as e:
        raise e


def coord(x, y, unit=1):
    """Get my size."""
    width, height = A4
    x, y = x * unit, height - y * unit
    return x, y


def get_style(style=False):
    """Get the style required."""
    try:
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name='Justify', alignment=TA_JUSTIFY, leading=22))
        styles.add(ParagraphStyle(
            name="Centered", alignment=TA_CENTER, leading=22))
        width, height = letter
        my_style = styles["Normal"]
        if style and style in styles:
            my_style = styles[style]
        return my_style
    except Exception:
        pass


def create_paragraph(c, text, x, y, style=False):
    """Method to do paragraphing."""
    try:
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name='Justify', alignment=TA_JUSTIFY, leading=22))
        styles.add(ParagraphStyle(
            name="Centered", alignment=TA_CENTER, leading=22))
        width, height = letter
        my_style = styles["Normal"]
        if style and style in styles:
            my_style = styles[style]
        p = Paragraph(text, style=my_style)
        p.wrapOn(c, width, height)
        p.drawOn(c, *coord(x, y, mm))
    except Exception:
        pass


def date_suffix(theday):
    """Show date suffix."""
    d = int(theday.strftime('%d'))
    st_rd = {1: 'st', 2: 'nd', 3: 'rd'}
    suffix = 'th' if 11 <= d <= 13 else st_rd.get(d % 10, 'th')
    return suffix


def simple_document(params, document_name='CPIMS', report_name='letter'):
    """Method to write forms."""
    from reportlab.pdfgen import canvas
    try:
        a4width, a4height = A4
        file_name = "%s/%s.pdf" % (MEDIA_ROOT, report_name)
        width, height = A4
        sst = getSampleStyleSheet()
        tgram = '.' * 34
        ref = '.' * 36
        if 'case_serial' in params:
            ref = str(params['case_serial'])

        logo = "%s/img/gok_logo.jpg" % (STATIC_ROOT)
        today = datetime.now()
        dt_suffix = date_suffix(today)
        suffix = '<sup>%s</sup>' % (dt_suffix)
        formatted_date = '{dt.day}{S} {dt:%B}, {dt.year}'.format(
            dt=today, S=suffix)

        p0 = Paragraph('''
            <font size=9>Telegram %s <br/>
            Telephone: 020 2077160<br/>
            When replying please quote<br/><br/>
            Ref. No. %s</font>
            ''' % (tgram, ref), sst["BodyText"])

        org_unit = params['org_unit']
        address = params['address']
        p1 = Paragraph('''
               <font size=9>%s<br/>
               %s<br/><br/>
               Date. %s</font>
               ''' % (org_unit, address, formatted_date), sst["BodyText"])
        im = Image(logo, 1.2 * inch, 1.0 * inch)
        data = [[p0, im, p1]]

        t = Table(data, style=[('ALIGN', (0, 0), (2, 0), 'CENTER'),
                               ('VALIGN', (0, 0), (2, 0), 'MIDDLE'), ])

        canvas = canvas.Canvas(file_name, pagesize=letter)
        fontname, fontsize = 'Times-Roman', 12
        # print canvas.getAvailableFonts()
        canvas.setLineWidth(.4)
        canvas.setFont(fontname, fontsize)
        # print canvas.getAvailableFonts()
        canvas.setAuthor('CPIMS')
        canvas.setTitle('CPIMS - %s' % (document_name))
        canvas.setSubject('Department of Children Services Official %s' % (
            document_name))
        story = []
        normal_style = get_style("Normal")
        center_style = get_style("Centered")
        heading = ('Ministry of East African Community (EAC), Labour and '
                   'Social Protection')
        sub_heading = 'DEPARTMENT OF CHILDREN SERVICES'
        ptext = '<font size=14><b>%s</b></font>' % heading.upper()
        story.append(Paragraph(ptext, center_style))
        ptext = '<font size=13><b>%s</b></font>' % sub_heading
        story.append(Paragraph(ptext, center_style))
        story.append(Spacer(1, 12))
        story.append(t)
        story.append(Spacer(1, 12))

        # Create return address
        ptext = ('<para align=center spaceb=3><font size=12>'
                 '<b>%s</b></font></para>') % document_name.upper()
        story.append(Paragraph(ptext, normal_style))
        # story.append(Spacer(1, 24))

        f = Frame(0.7 * inch, inch, 7.0 * inch, 9.2 * inch, showBoundary=0)
        f.addFromList(story, canvas)

        # create_paragraph(canvas, ptext, 10, 48, 'Centered')
        report_id = report_name.split('_')[0]
        fields = get_report_body(params, report=report_id)

        row_num = height - 300
        child_name = params['child_name']
        params = child_data(params)
        siblings = params['siblings']
        parents = params['guardians']
        p_count = len(parents)
        print "guardians count", p_count, parents
        p_counts = p_count if p_count >= 2 else 3

        for field in fields:
            s_field = ''
            field_value = fields[field]
            myfont = 'Times-Bold' if field.isupper() else 'Times-Roman'
            canvas.setFont(myfont, 12)
            twidth = stringWidth(field, fontname, fontsize)
            if '<sibling' in field:
                for i in range(1, 9):
                    sib_id = '%s.' % (i)
                    canvas.drawString(50, row_num + 3, sib_id)
                    if i == 1:
                        canvas.drawString(twidth, row_num + 3, child_name)
                    else:
                        if i in siblings:
                            sib_name = siblings[i]
                            canvas.drawString(twidth, row_num + 3, sib_name)
                        else:
                            canvas.line(twidth, row_num, width - 40, row_num)
                    row_num -= 20
                    row_num = paginate(row_num, canvas, myfont, height)
            if '<parents' in field:
                for i in range(1, p_counts):
                    sib_id = '%s.' % (i)
                    canvas.drawString(50, row_num + 3, sib_id)
                    if i in parents:
                        par_name = parents[i]
                        canvas.drawString(twidth, row_num + 3, par_name)
                    else:
                        canvas.line(twidth, row_num, width - 40, row_num)
                    row_num -= 20
                    row_num = paginate(row_num, canvas, myfont, height)
            if field_value or field_value == '':
                fd = unicode(str(field), 'utf-8')
                s_field = '.' if fd.isnumeric() else ':'
                if '<line' in field:
                    twidth = -20
                canvas.line(70 + twidth, row_num, width - 40, row_num)
                canvas.drawString(70 + twidth, row_num + 3, field_value)
            fid = s_field if '<blank' in field else '%s%s' % (field, s_field)
            chk_sib = '<sibling' not in field
            if '<line' not in field and chk_sib and '<parents' not in field:
                if field.startswith('<title_'):
                    ttext = re.sub(r'(?i)^<title_+\d+>', '', fid)
                    canvas.drawCentredString(width / 2, row_num + 3, ttext)
                else:
                    canvas.drawString(50, row_num + 3, fid)
                    # create_paragraph(canvas, 'ptext', 10, 48, 'Centered')
            row_num = paginate(row_num, canvas, myfont, height)
            row_num -= 20
        report_uid = uuid.uuid4()
        # Save in audit trail id,user_id,child_id,doc_id,org_id,datetime
        report_url = 'www.childprotection.go.ke/verify/%s' % (report_uid)
        qr_code = qr.QrCodeWidget(report_url)
        bounds = qr_code.getBounds()
        width = bounds[2] - bounds[0]
        height = bounds[3] - bounds[1]
        d = Drawing(45, 45, transform=[45. / width, 0, 0, 45. / height, 0, 0])
        d.add(qr_code)
        renderPDF.draw(d, canvas, a4width - (width), height)
        if report_id != 'DSUM':
            canvas.line(40, height, a4width - 40, height)

        canvas.save()
    except Exception as e:
        print 'error - %s' % (str(e))
        pass


def paginate(row_num, canvas, myfont, height):
    """Method to move to next page."""
    if row_num <= 80:
        canvas.showPage()
        canvas.setLineWidth(.4)
        canvas.setFont(myfont, 12)
        row_num = height - 100
    return row_num


def word_document(params, name):
    """Method to write to word - .docx."""
    try:
        from docx import Document
        from docx.shared import Inches

        document = Document()

        logo = "%s/img/gok_logo.jpg" % (STATIC_ROOT)

        document.add_heading('Document Title', level=0)

        p = document.add_paragraph('A plain paragraph having some ')
        p.add_run('bold').bold = True
        p.add_run(' and some ')
        p.add_run('italic.').italic = True

        document.add_paragraph(
            'first item in ordered list', style='List Number'
        )

        document.add_picture(logo, width=Inches(1.1))

        table = document.add_table(rows=1, cols=3)
        table.rows[0].style = "border:0;"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Qty'
        hdr_cells[1].text = 'Id'
        hdr_cells[2].text = 'Desc'

        document.add_page_break()
        document.save('demo.docx')

    except Exception as e:
        raise e


def draw_page(canvas, doc):
    """Method to format my pdfs."""
    title = "CPIMS"
    author = "CPIMS"
    canvas.setTitle(title)
    canvas.setSubject(title)
    canvas.setAuthor(author)
    canvas.setCreator(author)
    # footer = []
    # Put some data into the footer
    # Frame(2 * cm, 0, 17 * cm, 4 * cm).addFromList(footer, canvas)
    canvas.saveState()

    # Header
    canvas.drawString(0.5 * inch, 8 * inch, doc.fund_name)
    canvas.drawRightString(10.5 * inch, 8 * inch, doc.report_info)

    # Footers
    canvas.setFont("Helvetica", 8)
    canvas.drawString(0.5 * inch, 0.5 * inch, '')
    canvas.drawRightString(11.2 * inch, 0.5 * inch, 'Source : %s' % (doc.source))
    # canvas.drawRightString(10.5 * inch, 0.5 * inch, 'Page %d' % (doc.page))

    canvas.setFont("Helvetica", 240)
    # self.setFont("Helvetica", 8)
    canvas.setStrokeGray(0.90)
    canvas.setFillGray(0.90)
    # canvas.rotate(45)
    # canvas.drawCentredString(5.5 * inch, 3.25 * inch, doc.watermark)
    # logo = ImageReader('https://www.google.com/images/srpr/logo11w.png')
    # canvas.drawImage(logo, 10, 10, mask='auto')

    canvas.restoreState()


def get_sub_county_info(sub_county_ids, a_type='GDIS', icounty=None):
    """Get selected sub-counties info and attached counties."""
    area_ids, counties = {}, {}
    try:
        county_set = SetupGeography.objects.filter(
            area_type_id='GPRV', is_void=False)
        for county in county_set:
            counties[county.area_id] = county.area_name
        if icounty:
            sub_county_ids = SetupGeography.objects.filter(
                parent_area_id=icounty, area_type_id='GDIS',
                is_void=False).values_list('area_id', flat=True)
        if sub_county_ids:
            areas = SetupGeography.objects.filter(
                area_id__in=sub_county_ids, is_void=False)
            for area in areas:
                area_id = area.area_id
                area_type = area.area_type_id
                area_name = area.area_name
                parent_area_id = area.parent_area_id
                if parent_area_id in counties:
                    parent_area_id = counties[parent_area_id]
                if area_type == a_type:
                    area_ids[area_id] = {'county': parent_area_id,
                                         'sub_county_id': area_id,
                                         'sub_county': area_name}
    except Exception as e:
        print 'error getting sub-county ids - %s' % (str(e))
        return []
    else:
        return area_ids


def get_geo_locations(org_id, a_type='GDIS'):
    """Get specific Organisational units location based on org id."""
    ext_ids = []
    try:
        areas = RegOrgUnitGeography.objects.select_related().filter(
            org_unit_id=org_id, is_void=False)
        for area in areas:
            area_type = area.area.area_type_id
            area_name = area.area.area_name
            if area_type == a_type:
                ext_ids.append(area_name)
    except Exception as e:
        raise e
    else:
        return ext_ids


def get_period(report_type='M', month='', year='', period='F'):
    """
    Method to generate date ranges in preparation for the query.

    The report types include M, Q1, Q2, Q3 and Y
    period should be a calculated month range given an end date.
    """
    try:
        print 'TYPE', report_type, 'MON', month, 'YR', year, 'PERIOD', period,
        days = 30
        reports_qs = {'Q1': 9, 'Q2': 12, 'Q3': 3, 'Q4': 6}
        other_yr = ['Q3', 'Q4', 'Y']
        yr_add = 1 if report_type in other_yr else 0
        if period == 'C':
            yr_add = 0
            reports_qs = {'Q1': 3, 'Q2': 6, 'Q3': 9, 'Q4': 12}
        if not month and not year:
            today = datetime.now()
            year = today.strftime('%Y')
            month = today.strftime('%m')
        if report_type in reports_qs:
            days = 89 if report_type == 'Q3' else 90
            if period == 'C':
                days = 89 if report_type == 'Q1' else 90
            month = reports_qs[report_type]
        elif report_type == 'Y':
            # Days in year intentionally made < 365.25 by a day
            # for using timedelta replace day=1
            days = 364
            month = 12 if period == 'C' else 6
        elif report_type == 'U':
            days, month = 364, 9
        elif report_type == 'S':
            days, month = 180, 3
        year = int(year) + yr_add
        start_day_week, end_day = monthrange(int(year), int(month))
        end_date = '%s-%s-%s' % (end_day, month, year)
        end_date_obj = datetime.strptime(end_date, '%d-%m-%Y')
        # print (end_date_obj - timedelta(days=days))
        if days == 30:
            start_date_obj = end_date_obj.replace(day=1)
        else:
            sdate = (end_date_obj - timedelta(days=days))
            start_date_obj = sdate.replace(day=1)
        # DATIM Start
        datim_date = (end_date_obj - timedelta(days=89))
        datim_start = datim_date.replace(day=1)
        params = {}
        params['datim_start_date'] = datim_start
        params['end_date'] = end_date_obj
        params['start_date'] = start_date_obj
        # Period name for the report
        report_name = 'Quarterly' if report_type in reports_qs else 'Monthly'
        report_name = 'Annual' if report_type == 'Y' else report_name
        # Months names for the report
        mon_name = month_name[int(month)]
        sheet_name = '%s-%s' % (mon_name[:3], (str(year))[-2:])
        if report_type in reports_qs:
            s_mon = month_name[reports_qs[report_type] - 2][:3]
            mon_name = '%s-%s' % (s_mon, mon_name[:3])
            sheet_name = report_type.replace('Q', 'Qtr')
        mon_name = 'Jul-Jun' if report_type == 'Y' else mon_name
        if period == 'C':
            mon_name = 'Jan-Dec' if report_type == 'Y' else mon_name
        sheet_name = 'YEAR' if report_type == 'Y' else sheet_name
        # Start year
        s_year = year - yr_add
        year_name = '%s/%s' % (s_year, year) if report_type == 'Y' else year
        # Report label
        report_label = mon_name if report_type == 'M' else sheet_name
        params['period'] = report_name
        p_name = report_name[:-2] if report_name.endswith('ly') else 'Year'
        params['period_name'] = p_name
        params['month'] = mon_name
        params['year'] = year
        params['years'] = year_name
        params['sheet'] = sheet_name
        params['label'] = report_label
        print('BBBBBBBBBBBBBBBBBBBBBBB', params)
        return params
    except Exception as e:
        print 'error getting date - %s' % (str(e))
        raise e


def age_data(age, sex, vals, cat, summ=False):
    """Age set data - Upper value using range so its x-1."""
    ranges = [(0, 6), (6, 11), (11, 16), (16, 18)]
    for i, (lval, uval) in enumerate(ranges):
        if (lval <= age <= uval):
            val = i + 1
    if age >= 18:
        val = 5
    age_set = 'age_%d_%s' % (val, sex[1].lower())
    if summ:
        val = vals[cat][age_set]
        val.append(str(summ))
    else:
        val = vals[cat][age_set] + 1
    return age_set, val


def initial_values(cat, val, age, sex, intv=False):
    """Assign initial values - zeros."""
    age_sets = OrderedDict()
    for i in range(1, 6):
        age_sets['age_%s_m' % (i)] = 0
        age_sets['age_%s_f' % (i)] = 0
    if intv:
        cat = '%s_%s' % (intv, cat)
    val[cat] = age_sets
    age_set, dt = age_data(age, sex, val, cat)
    val[cat][age_set] = dt
    if intv:
        val[cat]['cat'] = intv


def sum_values(summary, cat, val, age, sex):
    """Get values."""
    age_sets = OrderedDict()
    for i in range(1, 6):
        age_sets['age_%s_m' % (i)] = []
        age_sets['age_%s_f' % (i)] = []
    val[summary] = age_sets
    age_set, dt = age_data(age, sex, val, summary, cat)
    val[summary][age_set] = dt


def data_from_results(datas, gid='D', intv=False):
    """Method to generate data set from result set.

    This has cat, sex, age
    """
    try:
        val, itv = {}, False
        for data in datas:
            cat = str(data['cat'])
            if intv:
                itv = cat
                cat = str(data['itv'])
            cid = str(data['cid'])
            kid = data['kid']
            age = data['age']
            sex = data['sex']
            if cat not in val:
                initial_values(cat, val, age, sex, itv)
            else:
                age_set, dt = age_data(age, sex, val, cat)
                val[cat][age_set] = dt
            # Summary values for unique cases and children
            if 'CASE' not in val:
                sum_values('CASE', cid, val, age, sex)
            else:
                age_set, dt = age_data(age, sex, val, 'CASE', cid)
                val['CASE'][age_set] = dt
            if gid == 'D':
                if 'CHILD' not in val:
                    sum_values('CHILD', kid, val, age, sex)
                else:
                    age_set, dt = age_data(age, sex, val, 'CHILD', kid)
                    val['CHILD'][age_set] = dt
        return val

    except Exception as e:
        raise e


def filter_org_unit(params, sub_county_case_ids):
    """Method to filter by org unit."""
    try:
        org_unit_tree = []
        org_unit_ids = []
        if 'org_unit_tree' in params:
            org_unit_tree = params['org_unit_tree']
            print 'tquery', org_unit_tree
            if len(org_unit_tree) > 1:
                org_unit_ids = org_unit_tree
        if 'org_unit' in params:
            org_unit = params['org_unit']
            if org_unit and len(org_unit_ids) == 0:
                print 'for one', org_unit
                org_unit_ids = [int(org_unit)]
        if org_unit_ids:
            sub_county_case_ids = sub_county_case_ids.filter(
                report_orgunit_id__in=org_unit_ids)
    except Exception as e:
        print "could not filter by OU - %s" % (str(e))
        return sub_county_case_ids
    else:
        return sub_county_case_ids


def get_data(params, report='CASE_LOAD'):
    """
    Method to do actual query from the db.

    For now lets do case data only.
    """
    try:
        data = []
        print 'Case Load params', params
        print('\n\n')
        cl_queryset = OVCCaseCategory.objects.all()
        cl_queryset = cl_queryset.filter(is_void=False).exclude(
            case_id__is_void=True)
        # Date ranges
        org_unit = params['org_unit'] if 'org_unit' in params else False
        sc = params['sub_county_id'] if 'sub_county_id' in params else False
        # Date filters
        # cl_queryset = filter_by_date(params, cl_queryset)
        case_ids = get_case_ids(params)
        cl_queryset = cl_queryset.filter(case_id_id__in=case_ids)
        # Now check for Region and org unit
        if sc:
            sub_county = params['sub_county_id']
            # Get all case ids for this sub_county
            sub_county_case_ids = OVCCaseGeo.objects.filter(
                report_subcounty_id__in=sub_county,
                is_void=False).values_list('case_id_id', flat=True)
        else:
            sub_county_case_ids = OVCCaseGeo.objects.filter(
                is_void=False).values_list('case_id_id', flat=True)
            if org_unit:
                print 'by ou', org_unit
                sub_county_case_ids = filter_org_unit(params,
                                                      sub_county_case_ids)
        cl_queryset = cl_queryset.filter(
            case_id_id__in=(sub_county_case_ids))

        for cl in cl_queryset:
            item = {}
            item['cat'] = cl.case_category
            item['sex'] = cl.person.sex_id
            item['age'] = cl.person.years
            # For generating summaries
            item['kid'] = cl.person.id
            item['cid'] = cl.case_id_id
            data.append(item)
        pending = get_pending(params)
        interventions = get_intervention(params, pending)
        raw_data = data_from_results(data)
        raw_pending = data_from_results(pending, gid='P')
        raw_interven = data_from_results(interventions, gid='I')
        intvs = data_from_results(interventions, gid='I', intv=True)
        raw_vals = {'data': raw_data, 'pending': raw_pending,
                    'interventions': raw_interven, 'itv': intvs}
        return raw_vals
    except Exception as e:
        print 'Get data error - %s' % (str(e))
        raise e


def data_category(data):
    """Method to show my category."""
    try:
        pass
    except Exception as e:
        raise e


def filter_by_date(params, queryset, field='date_of_event'):
    """Method to filter by dates provided."""
    try:
        if 'start_date' in params and 'end_date' in params:
            start_date = params['start_date']
            end_date = params['end_date']
            print "Filter by dates", str(start_date), str(end_date)
            kwargs = {'{0}__range'.format(field): (start_date, end_date)}
            queryset = queryset.filter(**kwargs)
    except Exception as e:
        print "Error applying date filter - %s" % (str(e))
        return queryset
    else:
        return queryset


def get_pending(params):
    """To get all pending within the same reporting period."""
    try:
        data = []
        pen_queryset = OVCCaseCategory.objects.all()
        pen_queryset = pen_queryset.filter(is_void=False)
        # Has interventions
        cats_itvs = OVCCaseEventServices.objects.filter(
            is_void=False).values_list('case_category_id', flat=True)
        pen_queryset = pen_queryset.exclude(case_category_id__in=cats_itvs)
        # Date ranges
        org_unit = params['org_unit'] if 'org_unit' in params else False
        sc = params['sub_county_id'] if 'sub_county_id' in params else False
        # Date filter
        # pen_queryset = filter_by_date(params, pen_queryset)
        case_ids = get_case_ids(params)
        pen_queryset = pen_queryset.filter(case_id_id__in=case_ids)
        # Missing filters
        if sc:
            sub_county = params['sub_county_id']
            # Get all case ids for this sub_county
            sub_county_case_ids = OVCCaseGeo.objects.filter(
                report_subcounty_id__in=sub_county,
                is_void=False).values_list('case_id_id', flat=True)
        else:
            sub_county_case_ids = OVCCaseGeo.objects.filter(
                is_void=False).values_list('case_id_id', flat=True)
            if org_unit:
                sub_county_case_ids = filter_org_unit(params,
                                                      sub_county_case_ids)
        pen_queryset = pen_queryset.filter(
            case_id_id__in=(sub_county_case_ids))
        for pen in pen_queryset:
            item = {}
            item['cat'] = pen.case_category
            item['sex'] = pen.person.sex_id
            item['age'] = pen.person.years
            # For generating summaries
            item['kid'] = pen.person.id
            item['cid'] = pen.case_id_id
            data.append(item)
        return data
    except Exception as e:
        print 'Get pending error - %s' % (str(e))
        raise e


def get_case_ids(params):
    """Method to filter cases that needs to be followed up."""
    try:
        period_qs = OVCCaseRecord.objects.filter(is_void=False)
        period_qs = filter_by_date(params, period_qs, 'date_case_opened')
        period_qs = period_qs.values_list('case_id', flat=True)
        print "case filters for ids by date"
        return period_qs
    except Exception as e:
        print "Error getting case -%s" % (str(e))
        return []


def get_intervention(params, pending=False):
    """To get all interventions within the same reporting period."""
    try:
        data = []
        itv_queryset = OVCCaseEventServices.objects.all()
        itv_queryset = itv_queryset.filter(is_void=False)
        # Start filtering , date_of_encounter_event__range
        org_unit = params['org_unit'] if 'org_unit' in params else False
        sc = params['sub_county_id'] if 'sub_county_id' in params else False
        if pending:
            pending_list = []
            for vals in pending:
                pending_list.append(vals['cid'])
            itv_queryset = itv_queryset.exclude(
                case_category__case_id_id__in=pending_list)
        case_ids = get_case_ids(params)
        itv_queryset = itv_queryset.filter(
            case_category__case_id_id__in=case_ids)
        # Missing filters
        if sc:
            sub_county = params['sub_county_id']
            # Get all case ids for this sub_county
            sub_county_case_ids = OVCCaseGeo.objects.filter(
                report_subcounty_id__in=sub_county,
                is_void=False).values_list('case_id_id', flat=True)
        else:
            sub_county_case_ids = OVCCaseGeo.objects.filter(
                is_void=False).values_list('case_id_id', flat=True)
            if org_unit:
                sub_county_case_ids = filter_org_unit(params,
                                                      sub_county_case_ids)
        itv_queryset = itv_queryset.filter(
            case_category__case_id__in=sub_county_case_ids)
        for res in itv_queryset:
            item = {}
            item['itv'] = res.service_provided
            item['cat'] = res.case_category.case_category
            item['kid'] = res.case_category.person_id
            item['age'] = res.case_category.person.years
            item['sex'] = res.case_category.person.sex_id
            item['cid'] = res.case_category.case_id_id
            data.append(item)
        return data
    except Exception as e:
        print 'Get intervention error - %s' % (str(e))
        raise e


def child_data(data):
    """Method to get child data from id."""
    try:
        parent_id = 0
        areas = {}
        areas['GPRV'] = 'child_county'
        areas['GDIS'] = 'child_sub_county'
        areas['GWRD'] = 'child_ward'
        # External ids
        extids = {}
        extids['ICOU'] = 'nationality'
        extids['IREL'] = 'religion'

        sibs, guards = {}, {}
        child_id = data['child_id']
        child_info = RegPerson.objects.get(pk=child_id)
        data['name'] = (str(child_info.full_name)).upper()
        data['age'] = (child_info.age).upper()
        # Get siblings
        siblings = RegPersonsSiblings.objects.filter(
            is_void=False, date_delinked=None, child_person_id=child_id)
        cnt = 2
        for sibling in siblings:
            sib_name = sibling.sibling_person.full_name
            sibs[cnt] = sib_name.upper()
            cnt += 1
        data['siblings'] = sibs
        # Get Guardians and parents
        guardians = RegPersonsGuardians.objects.filter(
            is_void=False, date_delinked=None, child_person_id=child_id)
        gcnt = 1
        for guardian in guardians:
            guard_name = guardian.guardian_person.full_name
            guards[gcnt] = guard_name.upper()
            gcnt += 1
        data['guardians'] = guards

        # Get child Geos
        geos = RegPersonsGeo.objects.filter(
            is_void=False, date_delinked=None, person_id=child_id,
            area_type='GLTL')
        for geo in geos:
            geo_type = geo.area.area_type_id
            if geo_type == 'GDIS':
                parent_id = int(geo.area.parent_area_id)
            area_name = geo.area.area_name
            data[areas[geo_type]] = area_name.upper()
        # Get county from list geo
        if 'child_county' not in data and parent_id != 0:
            county = SetupGeography.objects.get(area_id=parent_id)
            county_name = county.area_name
            data[areas['GPRV']] = county_name.upper()
        # Get child other details
        child_infos = RegPersonsExternalIds.objects.filter(
            is_void=False, person_id=child_id)
        ext_ids = {}
        for child_info in child_infos:
            child_ident = child_info.identifier_type_id
            child_ident_val = child_info.identifier
            ext_ids[child_ident] = child_ident_val
        for ffv in extids:
            ffv_key = extids[ffv]
            c_data = ext_ids[ffv] if ffv in ext_ids else 'N/A'
            data[ffv_key] = c_data
        return data
    except Exception as e:
        print 'Error getting child - %s' % (str(e))
        data['siblings'] = {}
        return data


def get_raw_data(params, data_type=1):
    """Method to filter which report user wants."""
    # class="table table-bordered"
    data = None
    raw_data = []
    td_zeros = '<td>0</td>' * 12
    tblanks = ['M', 'F'] * 6
    blanks = ['0'] * 13
    csvh = ['0 - 5 yrs', '', '6 - 10 yrs', '', '11 - 15 yrs', '',
            '16 - 17 yrs', '', '18+ yrs', '', 'Sub-Total', '', '']
    try:
        report_type = int(params['report_id'])
        if report_type == 6:
            data, raw_data = get_ovc_values(params)
        elif report_type == 5:
            data, raw_data = get_raw_values(params)
        elif report_type == 4:
            # Other values
            otherd, vls_ids = {}, {}
            idata, idatas = {}, {}
            otherd[1] = 'ALL OTHER DISEASES'
            otherd[2] = 'FIRST ATTENDANCES'
            otherd[3] = 'RE-ATTENDANCES'
            otherd[4] = 'REFERRALS IN'
            otherd[5] = 'REFERRALS OUT'
            vals = get_dict(field_name=['new_condition_id'])
            for vls in vals:
                vls_ids[vls] = vls
            if 'NCOD' in vals:
                del vals['NCOD']
            r_title = "{period_name}ly Health Report {unit_type}"
            dt = '<table class="table table-bordered"><thead>'
            dt += "<tr><th colspan='16'>%s" % (r_title)
            dt += '</th></tr>'
            dt += "<tr><th colspan='16'>{cci_si_name}</th></tr>"
            dt += case_load_header(report_type=3)
            dt += "</thead>"
            # Fetch raw data
            rdatas = get_institution_data(params, report_type)
            rdata = get_totals(rdatas['data'], vls_ids)
            if rdata:
                idata = write_row(rdata)
                idatas = write_row(rdata, is_raw=True)
            # Get totals
            itotals = col_totals(rdata)
            itotal = write_row([itotals])
            # Show the values
            total_h = 0
            diss_vals = {}
            hel_head = ['', 'Health Report'] + [''] * 13
            hel_title = ['', 'List of Diseases'] + tblanks + ['Total']
            if rdata:
                raw_data.append(hel_head)
                raw_data.append(['', ''] + csvh)
                raw_data.append(hel_title)
            cnt = 1
            other_items = {1: 'NCOD'}
            for val in vals:
                val_name = vals[val]
                val_data = diss_vals[val] if val in diss_vals else 0
                total_h += val_data
                dt += '<tr><td>%s</td><td>%s</td>' % (str(cnt) + '.', val_name)
                if val in idata:
                    dt += '%s' % (idata[val])
                else:
                    dt += '<td></td>%s<td>0</td></tr>' % (td_zeros)
                if val in idatas:
                    rd = idatas[val]
                    del rd[0:2]
                else:
                    rd = blanks
                if rdata:
                    raw_data.append([str(cnt) + '.', val_name] + rd)
                cnt += 1
            for oval in otherd:
                oval_name = otherd[oval]
                sval = other_items[oval] if oval in other_items else oval
                dt += '<tr><td>%s</td><td>%s</td>' % (str(cnt) + '.', oval_name)
                if sval in idata:
                    dt += '%s' % (idata[sval])
                else:
                    dt += '<td></td>%s<td>0</td></tr>' % (td_zeros)
                if sval in idatas:
                    rd = idatas[sval]
                    del rd[0:2]
                else:
                    rd = blanks
                if rdata:
                    raw_data.append([str(cnt) + '.', oval_name] + rd)
                cnt += 1
            if rdata:
                del itotals[1]
                raw_data.append([''] + itotals)
            dt += '<tr><td></td><td>Total</td>'
            dt += '%s' % (itotal['TOTAL'])
            dt += '<table>'
            data = dt
        elif report_type == 1:
            # KNBS List
            ids = {}
            ids['CSAB'] = 'Child offender'
            ids['CCIP'] = 'Children on the streets'
            ids['CSIC'] = 'Neglect'
            ids['CIDC'] = 'Orphaned Child'
            ids['CCIP'] = 'Children on the streets'
            ids['CDIS'] = 'Abandoned'
            ids['CHCP'] = 'Lost and found children'
            ids['CSDS'] = 'Drug and Substance Abuse'
            ids['CSNG'] = 'Physical abuse/violence'
            ids['CDSA'] = 'Abduction'
            ids['CCDF'] = 'Defilement'
            ids['CTRF'] = 'Child Labour'
            # Query just like case load
            all_datas = get_data(params)
            all_data = all_datas['data']
            knb_ids, rdata, rdatas = {}, {}, {}
            # Just send ids as ids for easier rendering later
            # Have to get all ids else errors
            case_categories = get_categories()
            for knb_id in ids:
                knb_ids[knb_id] = knb_id
            data = get_totals(all_data, case_categories)
            if data:
                rdata = write_row(data)
                rdatas = write_row(data, is_raw=True)
            rtotals = col_totals(data)
            rtotal = write_row([rtotals])
            rtitle = 'KNBS REPORT %s %s' % (
                params['month'], params['year'])
            # Just add title whether there is data or not
            knb_head = ['', rtitle.upper()] + [''] * 13
            knb_title = ['', 'Case Category'] + tblanks + ['Total']
            if data:
                raw_data.append(knb_head)
                raw_data.append(['', ''] + csvh)
                raw_data.append(knb_title)
            dt = '<table class="table table-bordered"><thead>'
            dt += '<tr><th colspan="16">%s</th></tr>' % (rtitle.upper())
            dt += case_load_header(report_type=4)
            dt += "</thead>"
            knbcnt = 1
            if data:
                for val in ids:
                    val_name = ids[val]
                    dt += '<tr><td>%s</td><td>%s</td>' % (knbcnt, val_name)
                    if val in rdata:
                        dt += '%s' % (rdata[val])
                    else:
                        dt += '<td></td>%s<td>0</td></tr>' % (td_zeros)
                    if val in rdatas:
                        rd = rdatas[val]
                        del rd[0:2]
                    else:
                        rd = blanks
                    raw_data.append([knbcnt, val_name] + rd)
                    knbcnt += 1
                raw_data.append(rtotals)
            dt += '<tr><td colspan="2"><b>Total</b></td>'
            dt += '%s' % (rtotal['TOTAL'])
            dt += '<table>'
            data = dt
        elif report_type == 3:
            discs = {'AEES': 'AEES'}
            dvals = {2: 'TANA', 4: 'TARR', 5: 'TRIN', 6: 'TARE'}
            rvals = {4: 'TARR', 5: 'TRIN', 6: 'TARE'}
            evals = {8: 'AEES', 9: 'AEAB', 10: 'TDER',
                     11: 'TDTR', 12: 'TDEX', 13: 'DTSI', 15: 'AEDE'}
            svals = {8: 'AEES', 9: 'AEAB', 10: 'TDER',
                     11: 'TDTR', 12: 'TDEX', 13: 'DTSI', 14: '14'}
            death_vals = {15: 'AEDE'}
            # Get all types of discharges
            discharges = get_dict(field_name=['discharge_type_id'])
            for disc in discharges:
                discs[disc] = disc
            # This is it
            pdatas = get_institution_data(params, report_type)
            devals = {}
            for dval in dvals:
                deq = dvals[dval]
                devals[deq] = deq
            pdata = get_totals(pdatas['data'], devals)
            odata = get_totals(pdatas['odata'], devals)
            ddata = get_totals(pdatas['ddata'], discs)
            edata = get_totals(pdatas['death'], death_vals)
            ids = {2: 'New Admissions',
                   3: 'Returnees',
                   4: ' - Relapse',
                   5: ' - Transfer in',
                   6: ' - Return after escape',
                   7: 'Exits',
                   8: ' - Escapee',
                   9: ' - Abducted',
                   10: ' - Early Release',
                   11: ' - Released on License',
                   12: ' - Released on Expiry of Order',
                   13: ' - Transfer to another Institution',
                   14: ' - Other exits',
                   15: 'Death'}
            r_title = "{cci_si_title} {period_name}ly Returns {unit_type}"
            dt = '<table class="table table-bordered"><thead>'
            dt += "<tr><th colspan='16'>%s" % (r_title)
            dt += "</th></tr><tr><th colspan='16'>{cci_si_name}</th></tr>"
            dt += case_load_header(report_type=2)
            dt += "</thead>"
            # This is it
            popdata, popdatas = {}, {}
            opdata, opdatas = {}, {}
            dopdata, dopdatas = {}, {}
            depdata, depdatas = {}, {}
            osdata = []
            if pdata:
                all_returnees = get_others(pdata, rvals, 3, True)
                pdata.append(all_returnees)
                popdata = write_row(pdata)
                popdatas = write_row(pdata, is_raw=True)
            # Old data
            if odata:
                p1 = col_totals(odata)
                osdata.append(p1)
                opdata = write_row(osdata)
                opdatas = write_row(osdata, is_raw=True)
            # Discharge data
            if ddata:
                all_other = get_others(ddata, evals, 14)
                ddata.append(all_other)
                all_exits = get_others(ddata, svals, 7, True)
                ddata.append(all_exits)
                dopdata = write_row(ddata)
                dopdatas = write_row(ddata, is_raw=True)
            # Deaths as a type of discharge
            if edata:
                depdata = write_row(edata)
                depdatas = write_row(edata, is_raw=True)
            # Just add title whether there is data or not
            pop_head = ['Institution Population'] + ['-'] * 14
            pop_title = ['Category', 'Sub-category'] + tblanks + ['Total']

            all_var = merge_two_dicts(popdata, dopdata)
            all_rvar = merge_two_dicts(popdatas, dopdatas)
            all_vars = merge_two_dicts(all_var, depdata)
            all_rvars = merge_two_dicts(all_rvar, depdatas)
            if pdata:
                raw_data.append(pop_head)
                raw_data.append(['', ''] + csvh)
                raw_data.append(pop_title)
            si_total = 0
            # All totals
            final_totals = get_final_totals(osdata, pdata, ddata)
            ptotal = write_row([final_totals])
            ptotal_raw = write_row([final_totals], is_raw=True)
            td_pad = '</td><td>'
            s_text = '<b>Total Children by End of Previous {period_name}</b>'
            dt += '<tr><td colspan="3">%s' % (s_text)
            if opdata:
                o_data = opdata['TOTAL'].replace('<td></td>', '')
                dt += o_data
                raw_ol = opdatas['TOTAL']
                del raw_ol[0:2]
                raw_data.append(['From previous period', ''] + raw_ol)
            else:
                dt += '</td>%s<td>0</td></tr>' % (td_zeros)
            ftotal = ptotal['TOTAL'].replace('<td></td>', '')
            for val in ids:
                vname = ids[val]
                v_name = vname.replace(' - ', td_pad)
                r_name = vname.replace(' - ', '')
                val_name = v_name + td_pad if '<td>' not in v_name else v_name
                vraw = [r_name, ''] if '<td>' not in v_name else ['', r_name]
                val_data = 0
                if val in dvals:
                    vd = dvals[val]
                elif val in evals:
                    vd = evals[val]
                else:
                    vd, val_data = str(val), 0
                dt += '<tr><td width="1px"></td><td>%s</td>' % (val_name)
                if vd in all_rvars:
                    rd = all_rvars[vd]
                    del rd[0:2]
                else:
                    rd = blanks
                if all_vars:
                    raw_data.append(vraw + rd)
                if vd in all_vars:
                    my_val = all_vars[vd].replace('<td></td>', '')
                    dt += '%s' % (my_val)
                else:
                    dt += '%s<td>0</td></tr>' % (td_zeros)
                si_total += val_data
            t_text = '<b>Total Children by End of Reporting {period_name}</b>'
            dt += '<tr><td colspan="3">%s</td>' % (t_text)
            dt += '%s' % (ftotal)
            dt += '</table>'
            if all_vars:
                raw_data.append(ptotal_raw['TOTAL'])
            data = dt
    except Exception as e:
        print('Error with raw data - %s' % (str(e)))
        raise e
    else:
        return data, raw_data


def get_age_bracket(age):
    """Method to do any age grouping."""
    val = 0
    agesets = ['', '0 - 5 yrs', '6 - 10 yrs', '11 - 15 yrs',
               '16 - 17 yrs', '18+ yrs']
    ranges = [(0, 6), (6, 11), (11, 16), (16, 18)]
    for i, (lval, uval) in enumerate(ranges):
        if (lval <= age <= uval):
            val = i + 1
    if age >= 18:
        val = 5
    bracket = agesets[val]
    return bracket


def get_ovc_values(params, data_type=1):
    """Get OVC report."""
    data = []
    ovc_filters = {'YEAR': 'ANNUAL', 'Qtr1': 'QUARTER1',
                   'Qtr3': 'QUARTER3', 'SemiAnnual': 'SEMIANNUAL'}
    try:
        print params
        rhead = ['OVCCount', 'Age', 'Agebracket', 'Domain', 'Gender',
                 'CBO', 'District', 'County', 'Ward']
        period = str(params['label'])
        year = params['year']
        rperiod = ovc_filters[period] if period in ovc_filters else 'ANNUAL'
        dt = '<table class="table table-bordered">'
        dt += '<thead><tr>'
        for hh in rhead:
            dt += '<th>%s</th>' % (hh)
        dt += '</tr></thead><tbody>'
        aggrs = OVCAggregate.objects.filter(
            project_year=year, reporting_period=rperiod)
        data.append(rhead)
        dt += '<tr><td colspan="9">Total Count %s</td></tr>' % (aggrs.count())
        for aggr in aggrs:
            a_1 = aggr.indicator_count
            a_2 = aggr.age
            ab = get_age_bracket(a_2)
            a_3 = aggr.indicator_name
            a_4 = aggr.gender
            a_5 = aggr.cbo
            a_6 = aggr.subcounty
            a_7 = aggr.county
            a_8 = aggr.ward
            data.append([a_1, a_2, ab, a_3, a_4, a_5, a_6, a_7, a_8])
        dt += '</tbody></table>'
    except Exception as e:
        raise e
    else:
        return dt, data


def get_raw_values(params, data_type=1):
    """Method to query summaries."""
    data = []
    try:
        dt = ''
        adhoc_type = int(params["adhoc_type"])
        if adhoc_type:
            print 'Run raw sql'
            adhoc_name = 'GOK_%s' % (adhoc_type)
            params['other_params'] = ''
            report_region = params['report_region']
            if report_region in [2, 3]:
                pms = params['sub_county_id']
                sc_list = [str(o_list) for o_list in pms]
                sc_ids = ','.join(sc_list)
                qs = 'AND cgeo.report_subcounty_id IN (%s)' % sc_ids
                params['other_params'] = qs
            if report_region == 4:
                org_unit_id = int(params['org_unit'])
                qs = 'AND cgeo.report_orgunit_id = %s' % (org_unit_id)
                params['other_params'] = qs
            rpt_name = REPORTS[adhoc_name]
            sql = QUERIES[rpt_name].format(**params)
            results, desc = run_sql_data(None, sql)
            # print datas, desc
            msg = 'Data Rendering Not available for this report.'
            dt += '<tr><td>Records (%s). %s</td></tr>' % (len(results), msg)
            titles = []
            if results:
                for res in results[0]:
                    titles.append(res)
            columns = [col.lower() for col in titles]
            data = [columns]
            print 'Results count - ', len(results)
            for res in results:
                vals = []
                for n, i in enumerate(titles):
                    val = res[i]
                    if type(val) is unicode:
                        val = val.encode('ascii', 'ignore').decode('ascii')
                    vals.append(val)
                data.append(vals)

    except Exception as e:
        raise e
    else:
        return dt, data


def get_institution_data(params, report_id=4):
    """Method to get all insititution data."""
    try:
        data = {}
        if report_id == 4:
            data = get_health_data(params)
        elif report_id == 3:
            data = get_population_data(params)
    except Exception as e:
        print "Error getting institution data - %s" % (str(e))
        return {}
    else:
        return data


def get_population_data(params):
    """
    Method to do actual query from the db.

    For now lets do institution values
    """
    try:
        data = []
        # Get institution filters - TNCI - charitable
        report_region = params['report_region']
        org_unit = params['org_unit']
        org_type = params['org_type']
        inst_type = params['institution_type']
        if org_type == 'ALL':
            si_list = ['TNRH', 'TNRB', 'TNRR', 'TNRS', 'TNAP']
            inst_list = ['TNRC'] if inst_type == 'TNCI' else si_list
        else:
            inst_list = [str(inst_type)]
        # Get all org units under this category
        if report_region == 4:
            org_list = [int(org_unit)]
        else:
            org_qs = RegOrgUnit.objects.filter(
                is_void=False, org_unit_type_id__in=inst_list)
            if report_region in [2, 3]:
                sub_ids = params['sub_county_id']
                orgs_geos = RegOrgUnitGeography.objects.filter(
                    is_void=False, area_id__in=sub_ids)
                orgs_geo = orgs_geos.values_list('org_unit_id', flat=True)
                uniq_orgs = list(set(orgs_geo))
                org_qs = org_qs.filter(org_unit_type_id__in=uniq_orgs)
            orgs_list = org_qs.values_list('id', flat=True)
            # This is a hack since Danet put this field as char and not int
            org_list = [str(o_list) for o_list in orgs_list]
        # Filter by date only first
        start_date = params['start_date']
        end_date = params['end_date']
        ip_queryset = OVCPlacement.objects.filter(
            residential_institution_name__in=org_list,
            is_void=False, admission_date__range=(start_date, end_date))
        place_ids = []
        for cl in ip_queryset:
            item = {}
            item['cat'] = cl.admission_type
            item['sex'] = cl.person.sex_id
            item['age'] = cl.person.years
            # For generating summaries
            item['kid'] = cl.person.id
            item['cid'] = cl.placement_id
            place_ids.append(cl.placement_id)
            data.append(item)
        # Last period population data
        dis_lists = OVCDischargeFollowUp.objects.filter(
            is_void=False, date_of_discharge__lt=start_date,
            placement_id_id__in=place_ids)
        dis_list = dis_lists.values_list('placement_id_id', flat=True)
        old_queryset = OVCPlacement.objects.filter(
            residential_institution_name__in=org_list,
            is_void=False, admission_date__lt=start_date)
        old_queryset = old_queryset.exclude(placement_id__in=dis_list)
        old_data = []
        for ol in old_queryset:
            oitem = {}
            oitem['cat'] = ol.admission_type
            oitem['sex'] = ol.person.sex_id
            oitem['age'] = ol.person.years
            oitem['kid'] = ol.person.id
            oitem['cid'] = ol.placement_id
            old_data.append(oitem)
        # All discharges
        dis_queryset = OVCDischargeFollowUp.objects.filter(
            is_void=False, date_of_discharge__range=(start_date, end_date),
            placement_id_id__in=place_ids)
        dis_data = []
        for dl in dis_queryset:
            ditem = {}
            ditem['cat'] = dl.type_of_discharge
            ditem['sex'] = dl.person.sex_id
            ditem['age'] = dl.person.years
            ditem['kid'] = dl.person.id
            ditem['cid'] = dl.placement_id_id
            dis_data.append(ditem)
        # Get all deaths within this period
        death_qs = OVCAdverseEventsFollowUp.objects.filter(
            is_void=False, adverse_condition_description='AEDE',
            adverse_event_date__range=(start_date, end_date),
            placement_id_id__in=place_ids)
        death_data = []
        for ds in death_qs:
            eitem = {}
            eitem['cat'] = ds.adverse_condition_description
            eitem['sex'] = ds.person.sex_id
            eitem['age'] = ds.person.years
            eitem['kid'] = ds.person.id
            eitem['cid'] = ds.placement_id_id
            print('ADVI', eitem)
            death_data.append(eitem)
        raw_data = data_from_results(data)
        raw_old = data_from_results(old_data)
        raw_dis = data_from_results(dis_data)
        raw_death = data_from_results(death_data)
        raw_vals = {'data': raw_data, 'odata': raw_old,
                    'ddata': raw_dis, 'death': raw_death}
        return raw_vals
    except Exception as e:
        print 'Get institution data error - %s' % (str(e))
        raise e


def get_health_data(params):
    """
    Method to do actual query from the db.

    For now lets do institution values
    """
    try:
        data = []
        # Get institution filters - TNCI - charitable
        report_region = params['report_region']
        org_unit = params['org_unit']
        org_type = params['org_type']
        inst_type = params['institution_type']
        if org_type == 'ALL':
            si_list = ['TNRH', 'TNRB', 'TNRR', 'TNRS', 'TNAP']
            inst_list = ['TNRC'] if inst_type == 'TNCI' else si_list
        else:
            inst_list = [str(inst_type)]
        # Get all org units under this category
        if report_region == 4:
            org_list = [int(org_unit)]
        else:
            org_qs = RegOrgUnit.objects.filter(
                is_void=False, org_unit_type_id__in=inst_list)
            if report_region in [2, 3]:
                print 'Filter further by county / sub-county'
                sub_ids = params['sub_county_id']
                orgs_geos = RegOrgUnitGeography.objects.filter(
                    is_void=False, area_id__in=sub_ids)
                orgs_geo = orgs_geos.values_list('org_unit_id', flat=True)
                uniq_orgs = list(set(orgs_geo))
                org_qs = org_qs.filter(org_unit_type_id__in=uniq_orgs)
            orgs_list = org_qs.values_list('id', flat=True)
            # This is a hack since Danet put this field as char
            org_list = [str(o_list) for o_list in orgs_list]
        # Filter by date only first
        start_date = params['start_date']
        end_date = params['end_date']
        '''
        ae_queryset = OVCAdverseEventsFollowUp.objects.filter(
            adverse_condition_description='AEMC',
            placement_id__residential_institution_name__in=org_list,
            is_void=False, adverse_event_date__range=(start_date, end_date))
        ae_lists = ae_queryset.values_list('adverse_condition_id', flat=True)
        '''
        ip_queryset = OVCAdverseEventsOtherFollowUp.objects.filter(
            is_void=False,
            adverse_condition_id__adverse_condition_description='AEMC')
        for cl in ip_queryset:
            item = {}
            person = cl.adverse_condition_id.person
            item['cat'] = cl.adverse_condition
            item['sex'] = person.sex_id
            item['age'] = person.years
            # For generating summaries
            item['kid'] = person.id
            item['cid'] = cl.adverse_condition
            data.append(item)
            print(data)
        raw_data = data_from_results(data)
        raw_vals = {'data': raw_data}
        return raw_vals
    except Exception as e:
        print 'Get institution data error - %s' % (str(e))
        raise e


def get_totals(all_data, categories, summ=False):
    """Method to get totals for case load."""
    data, cat = [], None
    cases, children = {}, {}
    try:
        print(categories)
        if not summ:
            if 'CASE' in all_data:
                cases = all_data['CASE']
                del all_data['CASE']
            if 'CHILD' in all_data:
                children = all_data['CHILD']
                del all_data['CHILD']
        for dt in all_data:
            print(dt)
            vls = all_data[dt]
            if 'cat' in vls:
                cat = vls['cat']
                del vls['cat']
            c_data = []
            sc, mc, fc = 0, 0, 0
            for vl in vls:
                sc += 1
                col_val = vls[vl]
                if dt == summ:
                    col_val = len(list(set(vls[vl])))
                c_data.append(col_val)
                if sc % 2:
                    mc += col_val
                else:
                    fc += col_val
            totals = [mc, fc, mc + fc]
            if dt == summ:
                vals = c_data + totals
                data.append(vals)
            else:
                if cat and '_' in cat:
                    dt, cat = cat.split('_')
                if '_' in dt:
                    dt, cat = dt.split('_')
                case_name = categories[dt]
                cat_name = categories[cat] if cat else ''
                vals = [case_name, cat_name] + c_data + totals
                data.append(vals)
        # Return the values
        if not summ:
            if cases:
                all_data['CASE'] = cases
            if children:
                all_data['CHILD'] = children
    except Exception as e:
        print('Totals error - %s' % (str(e)))
        pass
    else:
        return data


def col_totals(flist):
    """Method to get bottom totals."""
    ttls = ['TOTAL', '']
    try:
        sums = {}
        for i, fls in enumerate(flist):
            for x, fl in enumerate(fls):
                if x > 1:
                    if x in sums:
                        sums[x] = sums[x] + fl
                    else:
                        sums[x] = fl
        if sums:
            for tval in sums:
                tvalue = sums[tval]
                ttls.insert(tval, tvalue)
            ttotals = ttls
        else:
            ttotals = ttls + [0] * 13
        return ttotals
    except Exception as e:
        print 'Error get totals - %s' % (str(e))
        return ttls + [0] * 13


def get_final_totals(odata, pdata, ddata):
    """Method to now add up all totals."""
    try:
        my_totals = []
        # print 'OLD', odata
        data_old = get_others(odata, {1: 'TOTAL'}, 'OLD', True)
        my_totals.append(data_old)
        # print 'IN', pdata
        data_new = get_others(pdata, {1: 'TANA', 3: '3'}, 'IN', True)
        my_totals.append(data_new)
        # print 'OUT', ddata
        data_out = get_others(ddata, {1: 'AEDE', 7: '7'}, 'OUT', True)
        # Change this values to negatives
        new_data_out = prep_totals(data_out)
        my_totals.append(new_data_out)
        cal_sum = col_totals(my_totals)
        return cal_sum
    except Exception as e:
        raise e
    else:
        return []


def prep_totals(totals_list, exits=True):
    """Method to prepare list of addition."""
    try:
        new_list = []
        for x, fl in enumerate(totals_list):
            if exits:
                fll = fl * -1 if x > 1 else fl
            else:
                # Make this values positives
                fll = fl * -1 if x > 1 and fl < 0 else fl
            new_list.append(fll)
    except Exception:
        return totals_list
    else:
        return new_list


def get_others(ddata, evals, new_key=7, is_sum=False):
    """Method to do others."""
    try:
        flat_vars = []
        other_vars = []
        for ev in evals:
            flat_vars.append(evals[ev])
        for i, fls in enumerate(ddata):
            row_id = str(fls[0])
            if is_sum:
                if row_id in flat_vars:
                    other_vars.append(fls)
            else:
                if row_id not in flat_vars:
                    other_vars.append(fls)
        cal_sum = col_totals(other_vars)
        cal_sum[0] = new_key
        return cal_sum
    except Exception as e:
        print "error flattening other -%s" % (str(e))
        return []


def write_row(data, is_raw=False):
    """Method to write html rows given data."""
    try:
        row_dict = {}
        for val in data:
            row = [str(t) for t in val]
            cat_id = str(row[0])
            tmp_td = '<td>%s</td>' % (cat_id)
            table_string = "<td>" + \
                string.join(row, "</td><td>") + \
                "</td></tr>\n"
            if is_raw:
                row_dict[cat_id] = row
            else:
                row_dict[cat_id] = table_string.replace(tmp_td, '')
        return row_dict
    except Exception as e:
        print 'Row error - %s' % (str(e))
        return {}


def create_year_list(year_type='F', i_report=False):
    """Method to create year list for drop downs."""
    year_choices = []
    try:
        start_year = 2000 if i_report else 2015
        for yr in range(start_year, (datetime.now().year + 1)):
            yr_text = '%d/%d' % (yr, yr + 1) if year_type == 'F' else yr
            year_choices.append((yr, yr_text))
        return year_choices
    except Exception:
        return []


def get_categories(ids=True):
    """Method to get case categories."""
    try:
        case_categories = get_case_details(['case_category_id'])
        categories = {}
        for case_category in case_categories:
            case_id = case_category.item_id
            case_name = case_category.item_description
            if ids:
                categories[case_id] = case_id
            else:
                categories[case_id] = case_name
        return categories
    except Exception as e:
        print 'Error getting category - %s' % (str(e))
        return {}


def get_case_data(params):
    """Method to get case data."""
    try:
        case_id = params['case_id']
        case_info = {}
        # Case geo and serial
        case_reports = OVCCaseGeo.objects.filter(
            is_void=False, case_id_id=case_id)
        for case_report in case_reports:
            case_info['case_geo'] = case_report.report_subcounty.area_name
            case_info['case_serial'] = case_report.case_id.case_serial
        return case_info
    except Exception as e:
        print 'Get case error - %s' % (str(e))
        raise {}


def merge_two_dicts(x, y):
    """Given two dicts, merge them."""
    z = x.copy()
    z.update(y)
    return z


def org_unit_tree(org_unit_id):
    """Method to get all org units with all the children."""
    try:
        ou_ids = []
        qou_id = int(org_unit_id) if len(org_unit_id) > 0 else 0
        if qou_id not in ou_ids and qou_id > 0:
            ou_ids.append(qou_id)
        org_units = RegOrgUnit.objects.filter(
            parent_org_unit_id=qou_id, is_void=False)
        for org_unit in org_units:
            ou_id = org_unit.pk
            ou_ids.append(ou_id)
            # org_unit_tree(ou_id, ou_ids)
    except Exception as e:
        print "Error getting units tree - %s" % (str(e))
    else:
        return list(set(ou_ids))


def get_performance(request):
    """Method to get performance."""
    try:
        pds, ads = [], []
        punits, acases, cases = {}, {}, {}
        persons = RegPerson.objects.filter(is_void=False).values(
            'created_by__reg_person__first_name', 'created_by__reg_person_id',
            'created_by__reg_person__surname', 'created_by_id').annotate(
                person_count=Count('created_by_id')).order_by('-person_count')
        for pers in persons:
            pds.append(pers['created_by__reg_person_id'])
            ads.append(pers['created_by_id'])
        org_units = RegPersonsOrgUnits.objects.filter(
            is_void=False, person_id__in=pds)
        for unit in org_units:
            punits[unit.person_id] = unit.org_unit.org_unit_name
        # Get cases
        pcases = OVCCaseRecord.objects.filter(is_void=False).values(
            'created_by').annotate(
                case_count=Count('created_by')).order_by('-case_count')
        for case in pcases:
            acases[case['created_by']] = case['case_count']
        for pd in ads:
            if pd in acases:
                cases[pd] = acases[pd]
            else:
                cases[pd] = 0

    except Exception as e:
        print 'error with dashboard - %s' % (str(e))
    else:
        return persons, punits, cases


def get_performance_detail(request, user_id=0, params={}):
    """Method to get performance."""
    try:
        start_date = params['start_date']
        end_date = params['end_date']
        persons = RegPerson.objects.filter(
            is_void=False, created_by_id=user_id,
            created_at__range=(start_date, end_date)).values(
            'created_at', 'created_by_id').annotate(
                person_count=Count('created_at')).order_by('created_at')

        cases = OVCCaseRecord.objects.filter(
            is_void=False, created_by=user_id,
            timestamp_created__range=(start_date, end_date)).extra(
            select={'day': 'date( timestamp_created )'}).values(
            'day').annotate(case_count=Count('timestamp_created'))

        reports = OVCCaseRecord.objects.filter(
            is_void=False, created_by=user_id,
            date_case_opened__range=(start_date, end_date)).values(
            'date_case_opened').annotate(case_report=Count('date_case_opened'))

    except Exception as e:
        print 'error with performance - %s' % (str(e))
    else:
        return persons, cases, reports


def get_variables(request):
    """Method to prepare all the variables for reporting."""
    try:
        print request.POST
        dates = {v: k for k, v in enumerate(calendar.month_abbr)}
        sub_county_ids = request.POST.getlist('sub_county[]')
        sub_counties = request.POST.get('sub_county')
        county = request.POST.get('county')
        if not sub_county_ids:
            sub_county_ids = [sub_counties]
        report_type = request.POST.get('report_type')
        report_datim = request.POST.get('report_type_datim')
        rperiod = request.POST.get('report_period')
        rpt_years = request.POST.get('report_year')
        rpt_iyears = request.POST.get('report_years')
        report = request.POST.get('report_id', 0)
        report_unit = request.POST.get('org_unit')
        report_inst = request.POST.get('org_inst')
        org_unit_name = request.POST.get('org_unit_name')
        report_region = int(request.POST.get('report_region'))
        my_report_id = int(report)
        if my_report_id == 5:
            report_type = request.POST.get('report_type_other')
        # results = {'res': sub_county_ids}
        case_categories = get_case_details(['case_category_id'])
        institution_type = request.POST.get('institution_type')
        adhoc_id = request.POST.get('report_vars')
        org_type = request.POST.get('org_type')
        cluster = request.POST.get('cluster')
        report_ovc = request.POST.get('report_ovc')
        report_ovc_id = request.POST.get('rpt_ovc_id')
        if not report_type:
            report_type = request.POST.get('report_type_other')
        rpt_ovc = int(report_ovc) if report_ovc else 1
        rpt_ovc_id = int(report_ovc_id) if report_ovc_id else 1
        if rpt_ovc == 6:
            report_ovc_name = ORPTS[rpt_ovc_id]
        else:
            report_ovc_name = RPTS[rpt_ovc]
        report_name = report_ovc_name.title().replace(' ', '')
        categories = {}
        report_id = int(report) if report else 0
        if int(report_id) in [3, 4]:
            report_unit = report_inst
            rpt_years = rpt_iyears
        region_names = {1: 'National', 2: 'County',
                        3: 'Sub-county', 4: 'Organisation Unit',
                        5: 'Cluster'}
        is_fin = '/' in rpt_years if rpt_years else False
        report_year = rpt_years.split('/')[0] if is_fin else rpt_years
        for case_category in case_categories:
            case_id = case_category.item_id
            case_name = case_category.item_description
            categories[case_id] = case_name
        my_county = county if report_region == 2 else False
        if report_region == 1 or report_region == 4:
            sub_county_ids = []
        sub_counties = get_sub_county_info(
            sub_county_ids, icounty=my_county)
        variables = {'sub_county_id': [], 'sub_county': []}
        for sub_county in sub_counties:
            rep_var = sub_counties[sub_county]
            variables['county'] = rep_var['county']
            variables['sub_county_id'].append(rep_var['sub_county_id'])
            variables['sub_county'].append(rep_var['sub_county'])
        # Report variables
        if int(report_region) == 1:
            variables = {'county': 'National', 'sub_county': ['National']}
        variables['sub_county'] = ', '.join(variables['sub_county'])
        # Report variables
        variables['report_region'] = report_region
        variables['cluster'] = cluster if cluster else '0'
        today = datetime.now()
        # year = today.strftime('%Y')
        month = today.strftime('%m')
        # Other parameters
        if report_region == 4 or int(report_id) == 5:
            rep_unit = report_unit if report_unit else False
            variables['org_unit'] = rep_unit
        else:
            variables['org_unit'] = False
        variables['report_id'] = report_id
        rpd = rperiod[:3] if report_type == 'M' else rperiod
        ryr = rperiod[0] if report_type == 'Y' else 'F'
        if int(report_id) in [1]:
            ryr = 'C'
        year = int(report_year) + 1 if report_type == 'M' else report_year
        if report_type == 'Q':
            report_type = rperiod.replace('tr', '')
        # Pepfar us U for USG
        if rperiod == 'PEPFAR':
            report_type = 'U'
            if report_datim == 'S':
                report_type = 'S'
            year = int(report_year) + 1
        # Month value
        month = dates[rpd] if rpd in dates else ''
        # Handle calendars
        rpt_dt = request.POST.get('report_period')
        year = 2019
        if rpt_dt == 'Other':
            report_from_date = request.POST.get('report_from_date')
            report_to_date = request.POST.get('report_to_date')
            sdate = convert_date(report_from_date)
            edate = convert_date(report_to_date)
            period_params = {'start_date': sdate, 'end_date': edate}
            year = sdate.year
        else:
            period_params = get_period(
                report_type=report_type, year=year, month=month, period=ryr)
        report_variables = merge_two_dicts(variables, period_params)
        report_variables['org_unit_name'] = org_unit_name
        report_variables['institution_type'] = institution_type
        org_type_id = org_type if org_type else 'ALL'
        report_variables['org_type'] = org_type_id
        adhoc_type = int(adhoc_id) if adhoc_id else 0
        report_variables['adhoc_type'] = adhoc_type
        report_variables['report_ovc'] = report_ovc
        report_variables['report_ovc_name'] = report_name
        print('RID', report_id)
        greports = GRPTS[report_id] if report_id in GRPTS else 'CaseLoad'
        if report_id == 5:
            greports = ADHC[adhoc_type] if adhoc_type in ADHC else 'Ad hoc'
        report_variables['label'] = greports.replace(' ', '-')
        report_variables['year'] = year
        report_variables['month'] = month
        p_name = report_name[:-2] if report_name.endswith('ly') else 'Year'
        report_variables['period_name'] = p_name
        report_variables['cci_si_title'] = region_names[report_region]
        sc = variables['sub_county']
        is_opt = report_region == 4
        si_c = 'Sub-County: %s' % (sc) if report_region in [2, 3] else None
        si_n = 'National' if report_region == 1 else si_c
        si_name = 'Unit: %s' % org_unit_name if is_opt else si_n
        report_variables['cci_si_name'] = si_name
        for rval in request.POST:
            rvalue = request.POST.get(rval)
            if rval not in report_variables:
                report_variables[rval] = rvalue

        # More parameters
        inst_cats = {}
        inst_cats["TNCI"] = "Charitable Children Institution"
        inst_cats["TNSI"] = "Statutory Institution"
        # print('VARS 2', variables)

        case_institutions = get_case_details(
            ['si_unit_type_id', 'cci_unit_type_id'])
        for itsi in case_institutions:
            inst_id = itsi.item_id
            inst_name = itsi.item_description
            inst_cats[inst_id] = inst_name
        '''
        inst_type = institution_type if org_type_id == 'ALL' else org_type
        iname = inst_cats[inst_type] if inst_type in inst_cats else None

        icheck = iname is None and inst_type
        inst_type_name = inst_cats[inst_type] if icheck else iname
        '''
        unit_id = org_type_id if org_type_id != 'ALL' else institution_type
        check_region = unit_id and report_region != 4
        unit_type = 'for %s' % (inst_cats[unit_id]) if check_region else ''
        report_variables['unit_type'] = unit_type
        # print 'VARS', report_variables
        cbo_id = int(report_unit) if report_unit else 0
        if report_id == 5 and cluster:
            cbo_id = get_cbo_cluster(cluster)
        report_variables['cbos'] = cbo_id
        # print('RVARS', report_variables)
    except Exception as e:
        print 'error creating variables - %s' % (str(e))
        raise e
    else:
        return report_variables


def get_pivot_data(request, params={}):
    """Method to get pivot data."""
    try:
        field_names = ["case_category_id", "government_unit_type_id",
                       "ngo_unit_type_id", "cci_unit_type_id",
                       "si_unit_type_id", "committee_unit_type_id",
                       "adoption_unit_type_id"]
        print 'PERMS', params

        sq = ''
        report_region = params['report_region']
        if report_region in [2, 3]:
            sub_county_ids = params['sub_county_id']
            sub_ids = [str(o_list) for o_list in sub_county_ids]
            sids = ', '.join(sub_ids)
            sq = 'and report_subcounty_id in (%s)' % (sids)
        if report_region == 4:
            org_ids = params['org_unit']
            org_unit_ids = [int(org_ids)]
            sq = 'and report_orgunit_id in (%s)' % (org_ids)
        params['extras'] = sq
        sql = QUERIES['pivot_report'].format(**params)
        print('Start', datetime.now())
        data, cols = run_rawsql_data(None, sql, 1)
        df = pd.DataFrame(data)
        # print('df', df)
        # dv = df.values.tolist()
        print("Query done")
        # print(dv)
        columns = [col[0] for col in cols]
        results = df.values.tolist()
        # print(columns)
        fdata = [columns] + results
        # print(fdata)

        ndata = pd.DataFrame(results, columns=columns).fillna(0)

        ds = ndata.to_dict('records')
        '''
        for res in results:
            vals = []
            for n, i in enumerate(columns):
                val = res[i]
                vals.append(val)
            ds.append(vals)
        '''
        '''
        titles = []
        results, desc = run_sql_data(None, sql, 1)
        titles = []
        if results:
            for res in results[0]:
                titles.append(res)
        # columns = [col.lower() for col in titles]
        fdata = [titles] + [results]

        print('Results count - ', len(results))
        for res in results:
            vals = []
            for n, i in enumerate(titles):
                val = res[i]
                if type(val) is unicode:
                    val = val.encode('ascii', 'ignore').decode('ascii')
                vals.append(val)
            fdata.append(vals)
        # print(fdata)
        '''
        '''
        categories = get_dict(field_name=field_names)
        start_date = params['start_date']
        end_date = params['end_date']
        datas = OVCCaseCategory.objects.filter(
            case_id__date_case_opened__range=(start_date, end_date))
        gdatas = {}
        cids, cdatas = [], {}
        closes, cevents, cexcludes, oexcludes = {}, [], [], []
        case_ids = []
        for dt in datas:
            case_ids.append(dt.case_id_id)
        # Get the Geo details
        geos = OVCCaseGeo.objects.filter(case_id__in=case_ids)
        if report_region in [2, 3]:
            sub_county_ids = params['sub_county_id']
            geos_filters = geos.filter(report_subcounty_id__in=sub_county_ids)
            for geos_filter in geos_filters:
                cexcludes.append(geos_filter.case_id_id)
            datas = datas.filter(case_id__in=cexcludes)
        if report_region == 4:
            org_ids = params['org_unit']
            org_unit_ids = [int(org_ids)]
            orgs_filters = geos.filter(report_orgunit_id__in=org_unit_ids)
            for orgs_filter in orgs_filters:
                oexcludes.append(orgs_filter.case_id_id)
            datas = datas.filter(case_id__in=oexcludes)
        closures = OVCCaseEventClosure.objects.filter(
            case_event_id__case_id__in=case_ids)
        case_events = OVCCaseEvents.objects.filter(
            case_id__in=case_ids, date_of_event__range=(start_date, end_date))
        for closure in closures:
            closes[closure.case_event_id.case_id.case_id] = closure.case_outcome
        for case_event in case_events:
            cevents.append(case_event.case_id.case_id)
        for geo in geos:
            scname = geo.report_subcounty.area_name
            cname = geo.report_subcounty.parent_area_id
            otype = geo.report_orgunit.org_unit_type_id
            oname = geo.report_orgunit.org_unit_name
            cids.append(cname)
            gdatas[geo.case_id_id] = {'sub_county': scname, 'county': cname,
                                      'otype': otype, 'oname': oname}
        # Get the County names
        cids = list(set(cids))
        counties = SetupGeography.objects.filter(area_id__in=cids)
        for county in counties:
            cdatas[county.area_id] = county.area_name
        records = []
        ds = [['Category', 'Age', 'Unit Type', 'Sex', 'County', 'Sub County',
               'Organisation Unit', 'OVCCount', 'Case Status', 'Month',
               'Qtr', 'Case Date', 'System Date', 'Year']]
        for data in datas:
            case_id = data.case_id_id
            sex_id = data.person.sex_id
            age = data.person.years
            cat_id = data.case_category
            has_case = case_id in gdatas
            tid = gdatas[case_id]['otype'] if has_case else 'UK'
            county_id = gdatas[case_id]['county'] if has_case else 0
            county = cdatas[county_id] if county_id in cdatas else 'UK'
            scounty = gdatas[case_id]['sub_county'] if has_case else 'UK'
            category = categories[cat_id] if cat_id in categories else cat_id
            org_unit = gdatas[case_id]['oname'] if has_case else 'UK'
            unit_type = categories[tid] if tid in categories else tid
            status = ' Closed' if case_id in closes else 'Active'
            if status == 'Active':
                status = 'Active' if case_id in cevents else 'Pending'
            # county = data.case_id.ovccasegeo.occurrency_county
            sex = 'Male' if sex_id == 'SMAL' else 'Female'
            quotas = {1: 3, 2: 3, 3: 3, 4: 4, 5: 4, 6: 4,
                      7: 1, 8: 1, 9: 1, 10: 2, 11: 2, 12: 2}
            report_month = data.case_id.date_case_opened.strftime('%m-%b')
            rept_mon = int(data.case_id.date_case_opened.strftime('%m'))
            report_year = data.case_id.date_case_opened.strftime('%Y')
            case_date = data.case_id.date_case_opened.strftime('%d-%b-%Y')
            system_date = data.case_id.timestamp_created.strftime('%d-%b-%Y')
            quota = 'Qtr%s' % (quotas[rept_mon])
            report_quarter = ''
            record = {'Category': category, 'Age': age, 'Unit Type': unit_type,
                      'Sex': sex, 'County': county, 'Sub County': scounty,
                      'Organisation Unit': org_unit, 'OVCCount': 1,
                      'Status': status, 'Month': report_month, 'Qtr': quota,
                      'Case Date': case_date, 'System Date': system_date,
                      'Year': report_year}
            records.append(record)
            ds.append([category, age, unit_type, sex, county, scounty,
                       org_unit, 1, status, report_month, quota, case_date,
                       system_date, report_year])
        '''
        params['archive'] = True
        params['report_id'] = 100
        stime = '%s00' % (datetime.now().strftime('%Y%m%d%H'))
        file_name = '%s-CaseLoad.%s' % (request.user.id, stime)
        dx = int(params['report_excel']) if 'report_excel' in params else 0
        print(dx)
        excel_file = ''
        if dx == 1:
            excel_file, html = write_pd_csv(ndata, file_name, params)
        results = {'file_name': excel_file, 'records': ds, 'code': 0}
        print('End', datetime.now())
    except Exception as e:
        print 'error getting pivot data - %s' % (str(e))
        results = {'file_name': '', 'records': '', 'code': 9}
        return results
    else:
        return results


def get_geo_details(geo_ids):
    """Method to return geo ids dictionary."""
    try:
        areas = {}
        for county in range(1, 48):
            geo_ids.append(county)
        geos = SetupGeography.objects.filter(
            area_id__in=geo_ids)
        for geo in geos:
            area_id = geo.area_id
            area_name = geo.area_name
            vals = {'name': area_name}
            cid = geo.parent_area_id
            if cid:
                vals['county'] = cid
            areas[area_id] = vals
    except Exception as e:
        raise e
    else:
        return areas


def get_person_geodata(pers_ids):
    """Method to get cbo data RegPersonsGeo."""
    try:
        pers_data, sub_ids = {}, []
        pers_geos = RegPersonsGeo.objects.filter(
            is_void=False, person_id__in=pers_ids)
        sub_county_id = 0
        for pers in pers_geos:
            pers_id = pers.person_id
            area_type = pers.area.area_type_id
            area_name = pers.area.area_name
            pers_data[pers_id] = {}
            if area_type == 'GWRD':
                sub_county_id = pers.area.parent_area_id
                pers_data[pers_id]['ward'] = area_name
            if sub_county_id:
                pers_data[pers_id]['sub_county'] = sub_county_id
                sub_ids.append(sub_county_id)
        geo_ids = get_geo_details(sub_ids)
        for pdata in pers_data:
            sub_id = 0
            if 'sub_county' in pers_data[pdata]:
                sub_id = pers_data[pdata]['sub_county']
            if sub_id in geo_ids:
                sc_name = geo_ids[sub_id]['name']
                county_id = geo_ids[sub_id]['county']
                county_name = geo_ids[county_id]['name']
                pers_data[pdata]['sub_county'] = sc_name
                pers_data[pdata]['county'] = county_name
    except Exception as e:
        print 'error getting person geo data - %s' % (str(e))
        raise e
    else:
        return pers_data


def get_cbo_geodata(cbo_ids):
    """Method to get cbo data RegPersonsGeo."""
    try:
        cbo_data = {}
        cbos = RegOrgUnitGeography.objects.filter(
            is_void=False, org_unit_id__in=cbo_ids)
        for cbo in cbos:
            cbo_id = cbo.org_unit_id
            area_type = cbo.area.area_type_id
            area_name = cbo.area.area_name
            if area_type == 'GPRV':
                cbo_data[cbo_id] = {'county': area_name}
            if area_type == 'GDIS':
                cbo_data[cbo_id] = {'sub_county': area_name}
            if area_type == 'GWRD':
                cbo_data[cbo_id] = {'ward': area_name}
    except Exception as e:
        print 'error getting cbo geo data - %s' % (str(e))
        raise e
    else:
        return cbo_data


def get_domain_data(params):
    """Method to get data by domain."""
    try:
        datas = []
        domains = {}
        domains['olmis_shelter_service_id'] = 'Shelter and Care'
        domains['olmis_pss_service_id'] = 'Psychosocial Support'
        domains['olmis_protection_service_id'] = 'Protection'
        domains['olmis_hes_service_id'] = 'HouseHold Economic Strengthening'
        domains['olmis_health_service_id'] = 'Health and Nutrition'
        domains['olmis_education_service_id'] = 'Education'
        # sub domains
        field_names = ["olmis_shelter_service_id", "olmis_pss_service_id",
                       "olmis_protection_service_id", "olmis_hes_service_id",
                       "olmis_health_service_id", "olmis_education_service_id"]
        sub_domains = get_mapped(field_name=field_names)
        dmns = {}
        domain = ""
        for dmn in sub_domains:
            sdid = sub_domains[dmn]['id']
            dmns[str(dmn)] = sdid
            domain += "WHEN '%s' THEN '%s' " % (str(dmn), domains[sdid])
        end_date = params['end_date']
        start_date = params['start_date']
        cbo = params['org_unit']
        cbo_id = int(cbo) if cbo else 0
        rpt_id = params['report_region']
        cluster = params['cluster'] if 'cluster' in params else 0
        report_id = int(rpt_id) if rpt_id else 0
        if report_id == 5:
            cbo_id = get_cbo_cluster(cluster)
        params = {'cbos': cbo_id, 'start_date': start_date,
                  'end_date': end_date, 'domains': domain}
        sql = QUERIES['pepfar'].format(**params)
        datas, desc = run_sql_data(None, sql)
    except Exception as e:
        print 'error getting domain data - %s' % str(e)
    else:
        return datas


def get_kpi(kpis, reg_obj):
    """Method to get the parameter."""
    try:
        pass
    except Exception as e:
        raise e
    else:
        pass


def get_registration_data(kpis, params):
    """Get OVC registration data."""
    try:
        datas = []
        kpdts = [1, 2, 3, 4, 5, 7, 10, 11, 17, 24]
        end_date = params['end_date']
        start_date = params['start_date']
        ou = params['org_unit']
        cbo_id = int(ou) if ou else 0
        cbos = [cbo_id]
        # Handle clusters
        rpt_id = params['report_region']
        cluster = params['cluster'] if 'cluster' in params else 0
        report_id = int(rpt_id) if rpt_id else 0
        if report_id == 5:
            cbo_id = get_cbo_cluster(cluster)
            cbos = cbo_id.split(',')
        dt = date(start_date.year, start_date.month, start_date.day)
        regs = OVCRegistration.objects.filter(
            is_void=False, registration_date__lt=end_date,
            child_cbo_id__in=cbos)
        for reg in regs:
            age = reg.person.years
            sex = reg.person.sex_id
            cbo = reg.child_cbo.org_unit_name
            reg_date = reg.registration_date
            exit_date = reg.exit_date
            county = 1
            ward = 1
            kvar = 'Number of'
            gender = 'Female' if sex == 'SFEM' else 'Male'
            for kp in kpdts:
                if kp == 1:
                    kpi = kpis[kp] % (kvar)
                elif kp == 2 and reg_date > dt:
                    kpi = kpis[kp] % (kvar)
                elif kp == 3 and reg.is_active:
                    kpi = kpis[kp] % (kvar)
                elif kp == 4 and not reg.is_active:
                    kpi = kpis[kp] % (kvar)
                elif kp == 5 and not reg.is_active and exit_date > dt:
                    kpi = kpis[kp] % (kvar)
                elif kp == 7 and not reg.has_bcert:
                    kpi = kpis[kp] % (kvar)
                elif kp == 10 and reg.hiv_status == 'HSTP':
                    kpi = kpis[kp] % (kvar)
                elif kp == 11 and reg.hiv_status == 'HSTN':
                    kpi = kpis[kp] % (kvar)
                elif kp == 17 and reg.is_active and reg.hiv_status == 'HSTP':
                    kpi = kpis[kp] % (kvar)
                else:
                    kpi = None
                if kpi:
                    data = {'OVC Count': 1, 'Age': age,
                            'Gender': gender, 'CBO': cbo,
                            'County': county, 'Ward': ward,
                            'Performance Indicator': kpi}
                    datas.append(data)
    except Exception as e:
        raise e
    else:
        return datas


def format_data(rows, datas):
    """Get the data."""
    for row in rows:
        data = {'OVC Count': row[0], 'Age': int(row[5]),
                'Age Set': str(row[6]),
                'Gender': str(row[7]), 'CBO': row[2],
                'County': 'county', 'Ward': str(row[4]),
                'Services': str(row[8])}
        datas.append(data)


def get_services_data(servs, params):
    """Get OVC registration data."""
    try:
        datas = []
        start = time.clock()
        cbo = params['org_unit']
        cbo_id = int(cbo) if cbo else 0
        rpt_id = params['report_region']
        cluster = params['cluster'] if 'cluster' in params else 0
        report_id = int(rpt_id) if rpt_id else 0
        # cbos = [cbo_id]
        if report_id == 5:
            cbo_id = get_cbo_cluster(cluster)
            # cbos = cbo_id.split(',')
        # pids
        '''
        regs = OVCRegistration.objects.filter(
            is_void=False, registration_date__lt=end_date)
        if cbo_id:
            regs = regs.filter(child_cbo_id__in=cbos)
        # pids = regs.values_list('person_id', flat=True)
        # Results
        cbos = 'and ovc_registration.child_cbo_id in (%s)' % (cbo_id)
        if report_id == 1:
            cbos = ''
        '''
        params['cbos'] = cbo_id
        # HIVSTAT
        sql = QUERIES['datim'].format(**params)
        rows, desc = run_sql_data(None, sql)
        # Served
        sql1 = QUERIES['datim_1'].format(**params)
        sql3 = QUERIES['datim_3'].format(**params)
        sql4 = QUERIES['served'] % (sql1, sql3)
        rows1, desc1 = run_sql_data(None, sql4.replace(';', ''))
        # ART
        sql2 = QUERIES['datim_2'].format(**params)
        # print sql2
        rows2, desc2 = run_sql_data(None, sql2)
        #
        # format_data(rows, datas)
        # format_data(rows1, datas)
        # format_data(rows2, datas)
        datas = datas + rows + rows1 + rows2
        print 2, time.clock() - start
    except Exception as e:
        print 'datim error - %s' % (str(e))
        raise e
    else:
        return datas


def get_pivot_ovc(request, params={}):
    """Method to get OVC Pivot Data."""
    try:
        datas = []
        report_id = int(request.POST.get('report_ovc'))
        kpis = {}
        kpis[1] = '1.a %s OVCs Ever Registered'
        kpis[2] = '1.b %s New OVC Registrations within period'
        kpis[3] = '1.c %s Active OVC within period'
        kpis[4] = '2.a %s OVC Ever Exited from the program'
        kpis[5] = '2.b %s OVC Exited from the program withn period'
        kpis[6] = '2.c %s HouseHolds Exited from the program within period'
        kpis[7] = '3.a %s OVC without Birth Certificate at enrolment'
        kpis[8] = '3.b %s OVC served with Birth Certificate within period'
        kpis[9] = '3.c %s OVCs with Birth Certificates to date'
        kpis[10] = '3.d %s OVC served with Birth Certificate after enrolment'
        kpis[11] = ('3.e %s OVC 5yrs and below served with Birth Certificate '
                    'within period')
        kpis[12] = '4.a.(iii) %s OVC HIV Status not known at enrolment'
        kpis[13] = '4.a.(i) %s OVC HIV+ at enrolment'
        kpis[14] = '4.a.(ii) %s OVC HIV- at enrolment'
        kpis[15] = '4.b %s OVC Tested within period'
        kpis[16] = '4.c %s OVC Ever Tested for HIV'
        kpis[17] = '4.d %s OVC HIV+'
        kpis[18] = '4.e %s OVC HIV+ and NOT linked to Facilities'
        kpis[19] = '4.f %s OVC HIV+ and LINKED to Facilities'
        kpis[20] = '4.g %s ACTIVE OVC with known HIV status'
        kpis[21] = '4.h %s ACTIVE OVC HIV+'
        kpis[22] = '4.i %s ACTIVE OVC HIV+ and NOT linked to Facilities'
        kpis[23] = '4.j %s ACTIVE OVC HIV+ and LINKED to Facilities'
        kpis[24] = ('4.k %s ACTIVE OVC HIV+ and LINKED to Facilities '
                    'within period')
        kpis[25] = '4.l %s HIV+ OVC Tested within period'
        kpis[26] = '5.a %s ACTIVE HouseHolds/Caregivers within period'
        kpis[27] = '5.b %s Number of active CHVs within period'
        kpis[28] = '6.a %s OVC served with 1 or 2 services'
        kpis[29] = '6.b %s OVC served with 3 or more services'
        kpis[30] = '6.c %s OVC NOT served with any service'
        # served
        services = {}
        services[1] = 'a.OVC HIVSTAT'
        services[2] = 'b.OVC Served'
        services[3] = 'c.OVC Not Served'
        if report_id == 3:
            datas = get_registration_data(kpis, params)
        elif report_id == 2:
            datas = get_domain_data(params)
        elif report_id == 1:
            datas = get_services_data(services, params)
        else:
            datas, titles = get_sql_data(request, params)
    except Exception as e:
        print 'Error getting OVC pivot data - %s' % (str(e))
        return []
    else:
        return datas


def write_xls(response, data, titles=None):
    """Method to write excel."""
    try:
        wb = Workbook()
        ws = wb.active
        if titles:
            title = [i[0].upper() for i in titles]
            ws.append(title)
        for dt in data:
            ws.append(dt)
        # Save the file
        wb.save(response)
    except Exception as e:
        print "error creating excel - %s" % (str(e))
        raise e
    else:
        pass


def write_xlsm(csv_file, file_name, report_id=1):
    """Method to write excel."""
    try:
        print MEDIA_ROOT
        csv_file_name = '%s/%s.csv' % (MEDIA_ROOT, csv_file)
        excel_file = '%s/%s.xlsx' % (MEDIA_ROOT, file_name)
        s_name = RPTS[report_id] if report_id in RPTS else 1
        vba_file = '%s/%s/vbaProject.bin' % (DOC_ROOT, s_name)
        if os.path.isfile(vba_file):
            writer = pandas.ExcelWriter(excel_file, engine='xlsxwriter')
            data = pandas.read_csv(csv_file_name)
            data.to_excel(writer, sheet_name='Sheet1', index=False)
            workbook = writer.book
            xlsm_file = '%s/%s.xlsm' % (MEDIA_ROOT, file_name)
            workbook.add_worksheet('Sheet2')
            workbook.add_worksheet('Sheet3')
            workbook.filename = xlsm_file
            workbook.add_vba_project(vba_file)
            writer.save()
            writer.close()
            print 'Macros written - %s' % (xlsm_file)
        else:
            file_name = ""
            print 'No Macros Script - %s' % (vba_file)
    except Exception as e:
        print "error creating excel - %s" % (str(e))
        return ""
    else:
        return file_name


def write_csv(data, file_name, params):
    """Method to write csv given data."""
    try:
        html = ''
        csv_file = '%s/%s.csv' % (MEDIA_ROOT, file_name)
        with open(csv_file, 'wb') as csvfile:
            csvwriter = csv.writer(csvfile, delimiter=',', quotechar='"',
                                   quoting=csv.QUOTE_MINIMAL)
            csvwriter.writerows(data)
        mc = memcache.Client(['127.0.0.1:11211'], debug=0)
        print('PARAMS', params)
        dates = '%s' % (params['start_date'].strftime("%d, %b %Y"))
        dates += ' to %s' % (params['end_date'].strftime("%d, %b %Y"))
        mc.set(file_name, dates)
        # Save excel to flat file
        rid = params['report_id'] if 'report_id' in params else 1
        rheaders = False
        if rid == 5:
            rheaders = True
            rid = params['adhoc_type']
        report_id = int(rid) + 100
        s_name = RPTS[report_id] if report_id in RPTS else 100
        vba_file = '%s/%s/vbaProject.bin' % (DOC_ROOT, s_name)
        # print('VBA File', s_name, report_id)
        excel_file = ''
        if 'archive' in params:
            # print 'RID', file_name, s_name
            epoch_time = '%s00' % (datetime.now().strftime('%Y%m%d%H'))
            if file_name.startswith('tmp'):
                file_name = file_name.replace('tmp-', '')
                rnames = base64.urlsafe_b64decode(str(file_name))
                # print rnames
                report_details = rnames.split('_')
                s_name = '%s.%s' % (report_details[0], epoch_time)
                uid = report_details[-1]
                fname = '%s-%s' % (uid, s_name)
            else:
                fname = file_name
            df_new = pd.read_csv(csv_file)
            # Write to html
            html = df_new.to_html(
                index=False, header=rheaders, classes='table', na_rep='',
                float_format=None)
            excel_file = '%s.xlsx' % (fname)
            xlsx_file = '%s/xlsx/%s.xlsx' % (MEDIA_ROOT, fname)
            writer = pd.ExcelWriter(xlsx_file, engine='xlsxwriter')
            df_new.to_excel(writer, sheet_name='Sheet1', index=False)
            # This is it
            workbook = writer.book
            xlsm_file = '%s/xlsx/%s.xlsm' % (MEDIA_ROOT, fname)
            workbook.add_worksheet('Sheet2')
            workbook.add_worksheet('Sheet3')
            if os.path.isfile(vba_file):
                excel_file = excel_file.replace('xlsx', 'xlsm')
                workbook.filename = xlsm_file
                workbook.add_vba_project(vba_file)
            writer.save()
            writer.close()
            print 'Excel Files', xlsm_file, xlsx_file
    except Exception as e:
        print 'Error creating csv Results - %s' % (str(e))
        pass
    else:
        return excel_file, html


def write_pd_csv(df, file_name, params, is_csv=0):
    """Method to write csv given data."""
    try:
        html = ''
        if is_csv:
            print('Start csv')
            csv_file = '%s/%s.csv' % (MEDIA_ROOT, file_name)
            df.to_csv(csv_file, index=False, encoding='utf-8')
        # Now to excel
        report_id = 105
        s_name = RPTS[report_id] if report_id in RPTS else 100
        vba_file = '%s/%s/vbaProject.bin' % (DOC_ROOT, s_name)
        # Write to html
        print('Start excel')

        excel_file = '%s.xlsx' % (file_name)
        xlsx_file = '%s/xlsx/%s.xlsx' % (MEDIA_ROOT, file_name)

        df.to_excel(xlsx_file, index=False)
        print 'Excel Files', xlsx_file
    except Exception as e:
        print 'Error creating csv Results - %s' % (str(e))
        pass
    else:
        return excel_file, html


def get_sql_data(request, params):
    """Method to write data."""
    datas = []
    cbo = request.POST.get('org_unit')
    rpt_id = request.POST.get('report_region')
    report_ovc = request.POST.get('rpt_ovc_id')
    report_id = int(rpt_id) if rpt_id else 0
    rpt_ovc = int(report_ovc) if report_ovc else 1
    cluster = request.POST.get('cluster')
    cbo_id = int(cbo) if cbo else 0
    if report_id == 5:
        cbo_id = get_cbo_cluster(cluster)
    params['cbos'] = cbo_id
    print params
    df_rpt = REPORTS[1]
    qname = REPORTS[rpt_ovc] if rpt_ovc in REPORTS else df_rpt
    sql = QUERIES[qname]
    sql = sql.format(**params)
    print 'nnnnn'
    row, desc = run_sql_data(request, sql)
    data = datas + row
    qblank = '%s_blank' % (qname)
    if qblank in QUERIES:
        bsql = QUERIES[qblank]
        bsql = bsql.format(**params)
        brow, bdesc = run_sql_data(request, bsql)
        data = data + brow
    for i in range(1, 5):
        qs = '%s_%s' % (qname, str(i))
        qb = '%s_blank_%s' % (qname, str(i))
        if qs in QUERIES:
            sql = QUERIES[qs]
            sql = sql.format(**params)
            qrow, desc = run_sql_data(request, sql)
            data = data + qrow
        if qb in QUERIES:
            sql = QUERIES[qb]
            sql = sql.format(**params)
            brow, desc = run_sql_data(request, sql)
            data = data + brow
    return data, desc


def dictfetchall(cursor, cs=0):
    "Return all rows from a cursor as a dict"
    column = [col[0] for col in cursor.description]
    columns = [col.upper() for col in column]
    if cs:
        columns = column
    return [
        collections.OrderedDict(zip(columns, row))
        for row in cursor.fetchall()
    ]


def run_sql_data(request, sql, cs=0):
    """
    Method to handle Database connections
    User Reporting DB for reporting but
    Fallback to Transaction DB if any error
    """
    try:
        db_inst = 'default'
        dbinstance = connections[db_inst]
        print 'Query Reporting database .....'
        with dbinstance.cursor() as cursor:
            cursor.execute(sql)
            desc = cursor.description
            rows = dictfetchall(cursor, cs)
    except Exception as e:
        print 'Defaulting to Transaction DB - %s' % (str(e))
        with connection.cursor() as cursor:
            cursor.execute(sql)
            desc = cursor.description
            rows = dictfetchall(cursor, cs)
        return rows, desc
    else:
        return rows, desc


def run_rawsql_data(request, sql, cs=0):
    """
    Method to handle Database connections
    User Reporting DB for reporting but
    Fallback to Transaction DB if any error
    """
    try:
        db_inst = 'default'
        dbinstance = connections[db_inst]
        print 'Query Reporting database .....'
        with dbinstance.cursor() as cursor:
            cursor.execute(sql)
            # pd.read_sql_query(sql, cursor)
            desc = cursor.description
            data = cursor.fetchall()
    except Exception as e:
        print 'Defaulting to Transaction DB - %s' % (str(e))
        with connection.cursor() as cursor:
            cursor.execute(sql)
            desc = cursor.description
            data = cursor.fetchall()
        return data, desc
    else:
        return data, desc


def get_cbo_cluster(cluster_id):
    """Method to return cbo ids from clusters."""
    try:
        cbo_list = []
        my_list = OVCClusterCBO.objects.filter(
            is_void=False,
            cluster_id=cluster_id).values_list('cbo_id', flat=True)
        for a_list in my_list:
            cbo_list.append(str(a_list))
        cbos = ', '.join(cbo_list)
    except Exception as e:
        error = 'Error getting cluster cbo - %s' % (str(e))
        print error
        return 0
    else:
        return cbos


def get_cluster(request, id=0):
    """Method to return clusters."""
    try:
        cbo_id = request.session.get('ou_primary', 0)
        # Get all person ids attached to same org units as this person
        person_ids = RegPersonsOrgUnits.objects.filter(
            org_unit_id=cbo_id, is_void=False, date_delinked=None).values_list(
            'person_id', flat=True)
        # Find all user ids for this person
        user_ids = AppUser.objects.filter(
            reg_person_id__in=person_ids).values_list('id', flat=True)
        # Query all Clusters attached to this user ids
        cbos = OVCClusterCBO.objects.select_related().filter(
            cluster__created_by_id__in=user_ids)
        cs, clusters = {}, []
        for cbo in cbos:
            cbo_cluster = cbo.cluster.cluster_name
            cbo_name = cbo.cbo.org_unit_name
            cbo_id = cbo.cbo_id
            cluster_id = cbo.cluster_id
            cluster_date = cbo.cluster.created_at
            cluster_create = cbo.cluster.created_by
            if cluster_id not in cs:
                cs[cluster_id] = {'cluster_name': cbo_cluster,
                                  'created_at': cluster_date,
                                  'created_by': cluster_create,
                                  'cbos': [cbo_name], 'id': str(cluster_id),
                                  'cbo_ids': [cbo_id]}
            else:
                cs[cluster_id]['cbos'].append(cbo_name)
                cs[cluster_id]['cbo_ids'].append(cbo_id)
        for ci in cs:
            clusters.append(cs[ci])
    except Exception as e:
        error = 'Error getting cluster - %s' % (str(e))
        print error
        return []
    else:
        return clusters


def edit_cluster(request, cluster_id):
    """Cluster details."""
    try:
        status = 0
        user_id = request.user.id
        cluster_name = request.POST.get('cluster')
        cbos = request.POST.getlist('cbo')
        cboids = [int(cbo) for cbo in cbos]
        if cluster_id:
            dcl = OVCCluster.objects.get(id=cluster_id)
            dcl.delete()
            return 0
        if len(cbos) == 1:
            status = 1
        else:
            orgs = RegOrgUnit.objects.filter(
                is_void=False, id__in=cbos)
            for org in orgs:
                if org.parent_org_unit_id == 2:
                    status = 2
            if status != 2:
                # Is not an LIP
                clusters = get_cluster(request, id=0)
                for cs in clusters:
                    cbo_ids = cs['cbo_ids']
                    cboc = collections.Counter(cbo_ids)
                    if collections.Counter(cboids) == cboc:
                        status = 3
                if status != 3:
                    # Create Cluster
                    newc = OVCCluster(cluster_name=cluster_name,
                                      created_by_id=user_id)
                    newc.save()
                    # Attach CBOs to clusters
                    for cboid in cboids:
                        OVCClusterCBO(cluster_id=newc.pk,
                                      cbo_id=cboid).save()
    except Exception as e:
        print 'error - %s' % (str(e))
        return 9
    else:
        return status


def get_clusters(user, default_txt=False):
    """Method to return clusters."""
    initial_list = {'': default_txt} if default_txt else {}
    all_list = collections.OrderedDict(initial_list)
    try:
        cbo_ids = RegPersonsOrgUnits.objects.filter(
            person_id=user.reg_person_id, is_void=False,
            date_delinked=None).values_list('org_unit_id', flat=True)
        person_ids = RegPersonsOrgUnits.objects.filter(
            org_unit_id__in=cbo_ids, is_void=False,
            date_delinked=None).values_list('person_id', flat=True)
        # Find all user ids for this person
        user_ids = AppUser.objects.filter(
            reg_person_id__in=person_ids).values_list('id', flat=True)
        # Query all Clusters attached to this user ids
        if user.is_superuser:
            my_list = OVCCluster.objects.filter(
                is_void=False).order_by('cluster_name')
        else:
            my_list = OVCCluster.objects.filter(
                is_void=False,
                created_by_id__in=user_ids).order_by('cluster_name')
        for a_list in my_list:
            unit_names = a_list.cluster_name
            all_list[a_list.id] = unit_names
    except Exception as e:
        error = 'Error getting list - %s' % (str(e))
        print error
        return ()
    else:
        return all_list.items


def csvxls_data(request, f):
    """Method to convert csv to xlsx."""
    try:
        data = []
        csv_file = '%s/tmp-%s.csv' % (MEDIA_ROOT, f)
        with open(csv_file, 'rb') as csvfile:
            rows = csv.reader(csvfile, delimiter=',', quotechar='"')
            for row in rows:
                data.append(row)
    except Exception as e:
        print 'error - %s' % (str(e))
        return [], []
    else:
        return data, []


def create_pivot(df):
    """create pivot."""
    try:
        dt = pd.pivot_table(df, index=['name', 'active'],
                            columns=['services', 'agerange', 'gender'],
                            values=['ovccount'], margins=False,
                            aggfunc=[np.sum], fill_value=0)
    except Exception as e:
        raise e
    else:
        return dt


def create_sheet(df, writer, sheet):
    """Method to write sheets."""
    try:
        p_len = len(df.values)
        ws = writer.sheets[sheet]
        # writer.sheet = sheet
        ws.style = 'Normal'
        ecell = 'AS6:AS%s' % (5 + p_len)
        ccell = 'BX:ARX'.replace('X', str(6 + p_len))
        df.to_excel(writer, sheet_name=sheet, startrow=4,
                    startcol=0, header=False)
        # Totals
        mygreen = colors.Color(rgb='32CD32')
        gfill = PatternFill(start_color=mygreen, fill_type="solid")
        cnt = 6
        for cells in ws[ecell]:
            for cell in cells:
                cform = '=SUM(CX:ADX)/BX'.replace('X', str(cnt))
                cell.value = cform
                cell.fill = gfill
                cell.number_format = '0.00%'
                cnt += 1
        tcell = 'AX'.replace('X', str(6 + p_len))
        ws[tcell] = 'Grand Totals'
        ws[tcell].fill = gfill
        clt = 1
        for cells in ws[ccell]:
            for cell in cells:
                clt += 1
                clm = openpyxl.utils.get_column_letter(clt)
                if p_len > 0:
                    cell.value = '=SUM(%s6:%s%s)' % (clm, clm, str(5 + p_len))
                # cell.font = Font(color=colors.GREEN)
                cell.fill = gfill
        # Totals
        gcell = 'ASX'.replace('X', str(6 + p_len))
        cform = '=SUM(CX:ADX)/BX'.replace('X', str(cnt))
        if p_len > 0:
            ws[gcell] = cform
        ws[gcell].fill = gfill
        ws[gcell].number_format = '0.00%'
    except Exception as e:
        raise e
    else:
        pass


def create_pepfar(request, response, cfile):
    """Method to create PEPFAR."""
    try:
        csv_file = '%s/tmp-%s.csv' % (MEDIA_ROOT, cfile)
        print csv_file
        df = pd.read_csv(csv_file, sep=',')
        df_level = df[['level', 'name', 'active', 'services', 'agerange',
                       'gender', 'ovccount']]
        # To writing
        report_name = 'Pepfar'
        tmpt_file = '%s/%s.xltm' % (DOC_ROOT, report_name)
        writer = pd.ExcelWriter(response, engine='openpyxl')
        wb = openpyxl.load_workbook(tmpt_file)
        writer.book = wb
        writer.sheets = dict((ws.title, ws) for ws in wb.worksheets)
        for sheet in wb.sheetnames:
            # ws = wb[sheet]
            if sheet.startswith('CBO'):
                dfl = df_level[df.level == 'CBO']
                dataframe = create_pivot(dfl)
            elif sheet.startswith('Ward'):
                dfl = df_level[df.level == 'Ward']
                dataframe = create_pivot(dfl)
            else:
                dfl = df_level[df.level == 'County']
                dataframe = create_pivot(dfl)
            create_sheet(dataframe, writer, sheet)
        wb.save(response)
    except Exception as e:
        raise e
    else:
        pass


def get_dashboard_summary(request, report_id, category_id=0):
    """Method to get dashboard data."""
    try:
        datas = []
        cbo_id = request.session.get('ou_primary', 0)
        if request.user.is_superuser:
            cbo_id = 2
        print cbo_id, category_id, report_id
        if category_id == 1:
            # This is for USG
            if report_id == 4:
                orgs = RegOrgUnit.objects.filter(
                    parent_org_unit_id=cbo_id).order_by('-created_at')
                cnt = 0
                for org in orgs:
                    cnt += 1
                    val = {"id": cnt, "name": org.org_unit_name,
                           "create_date": org.created_at,
                           "action": org.org_unit_id_vis}
                    datas.append(val)
        else:
            if report_id == 14:
                orgs = RegOrgUnit.objects.filter(
                    parent_org_unit_id=cbo_id).order_by('-created_at')
                cnt = 0
                for org in orgs:
                    cnt += 1
                    val = {"id": cnt, "name": org.org_unit_name,
                           "create_date": org.created_at,
                           "action": org.org_unit_id_vis}
                    datas.append(val)
    except Exception as e:
        print 'error getting dashboard data - %s' % str(e)
        return []
    else:
        return datas


def get_org_unit(request, org_unit_id):
    """Method to get org unit."""
    try:
        org_unit = RegOrgUnit.objects.get(org_unit_id_vis=org_unit_id)
    except Exception:
        return org_unit_id
    else:
        return org_unit.org_unit_name


def get_styles():
    """Method to get styles."""
    try:
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name='Center', alignment=TA_CENTER))
        styles.add(ParagraphStyle(name='Right', alignment=TA_RIGHT))
        styles.add(ParagraphStyle(name='Left', alignment=TA_LEFT))
        styles.add(ParagraphStyle(
            name='Line_Data', alignment=TA_LEFT, fontSize=8,
            leading=11, fontColor='#FFFFFF'))
        styles.add(ParagraphStyle(
            name='Line_Data_Small', alignment=TA_LEFT,
            fontSize=7, leading=8))
        styles.add(ParagraphStyle(
            name='Line_Data_Large', alignment=TA_LEFT,
            fontSize=12, leading=15))
        styles.add(ParagraphStyle(
            name='Line_Data_Largest', alignment=TA_LEFT,
            fontSize=14, leading=15))
        styles.add(ParagraphStyle(
            name='Line_Label', font='Helvetica-Bold',
            fontSize=7, leading=6, alignment=TA_LEFT))
        styles.add(ParagraphStyle(
            name='Line_Label1', font='Helvetica-Bold', fontSize=7,
            leading=6, alignment=TA_RIGHT))
        styles.add(ParagraphStyle(
            name='Line_Label_Center', font='Helvetica-Bold',
            fontSize=12, alignment=TA_CENTER))
    except Exception as e:
        raise e
    else:
        return styles


def get_header(element, report_name, region, dates, styles):
    """Method to get header."""
    try:
        datenow = datetime.now()
        tarehe = datenow.strftime("%d, %b %Y %I:%M %p")
        url = 'https://childprotection.go.ke'
        # Handle headers
        address = '<b>MINISTRY OF LABOUR AND SOCIAL PROTECTION'
        address += "<br />STATE DEPARTMENT FOR SOCIAL PROTECTION"
        address += "<br />DEPARTMENT OF CHILDREN'S SERVICES</b>"
        report_number = '%s\n%s CPIMS Report\n%s' % (url, report_name, tarehe)
        bar_code = BarCode(value='%s' % (report_number))
        # Logo
        logo = "%s/img/logo_gok.png" % (STATIC_ROOT)
        sd = Image(logo)
        sd.drawHeight = 0.6 * inch
        sd.drawWidth = 0.7 * inch
        srg = 'CPIMS %s Report' % (region)
        data1 = [[sd, Paragraph(address, styles["Line_Data_Large"]),
                  Paragraph(srg, styles["Line_Data_Large"]), bar_code],
                 ['', '', Paragraph(dates, styles["Line_Label"]), '']]
        tt = Table(
            data1, colWidths=(2.1 * cm, None, 5.9 * cm, 3.4 * cm,),
            rowHeights=(1.2 * cm, .5 * cm,))

        tt.setStyle(TableStyle([
            ('INNERGRID', (2, 0), (2, 1), 0.25, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        discl = '<b>NOTE: This is a computer generated report <br /> '
        discl += ' as at %s.</b>' % (tarehe)
        element.append(tt)
        element.append(Spacer(0.1 * cm, .2 * cm))
        data1 = [[Paragraph('Report<br />Name:', styles["Line_Label"]),
                  Paragraph(report_name, styles["Line_Data_Largest"]),
                  Paragraph(discl, styles["Line_Data_Small"]), ]]

        t0 = Table(data1, colWidths=(2 * cm, None, 6.1 * cm,))
        t0.setStyle(TableStyle([
            ('INNERGRID', (1, 0), (-1, -1), 0.25, colors.black),
            ('BOX', (0, 0), (-1, -1), 0.25, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))

        element.append(t0)
    except Exception as e:
        raise e
    else:
        pass


def write_pdf(request, response, file_name):
    """Method to write to pdf from pandas."""
    try:

        # datas = pd.DataFrame(data)
        rparams = file_name.split('_')
        rid = int(rparams[3])
        region = rparams[0].replace('-', ' ')
        report_name = rparams[1].replace('-', ' ')
        ou = None
        if 'OU ' in region:
            ou_id = region.replace('OU ', '')
            ou_name = get_org_unit(request, ou_id)
            report_name = "%s (%s)" % (report_name, ou_name)
        styles = get_styles()
        element = []
        # Get headers for the report
        mc = memcache.Client(['127.0.0.1:11211'], debug=0)
        dates = mc.get(file_name.replace('.pdf', ''))
        get_header(element, report_name, region, dates, styles)
        csv_file = '%s/%s' % (MEDIA_ROOT, file_name)
        df = pd.read_csv(csv_file.replace('.pdf', '.csv'), na_filter=False)
        col_size = len(df.columns)
        dt_size = len(df.index)
        ds = dt_size - 1
        style = TableStyle(
            [('INNERGRID', (0, 0), (-1, -1), 0.25, colors.black),
             ('BOX', (0, 0), (-1, -1), 0.5, colors.black),
             ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
             ('ALIGN', (2, 2), (-1, -1), 'RIGHT'),
             ('FONTNAME', (0, ds), (-1, -1), 'Helvetica-Bold'),
             ('FONTNAME', (col_size - 1, 0), (-1, -1), 'Helvetica-Bold'),
             ('BACKGROUND', (0, 0), (-1, 0), '#89CFF0')])
        # Configure columns
        cols0 = {1: 7.9 * cm, 3: 5.7 * cm, 4: 8.55 * cm}
        cols1 = {1: 1.5 * cm, 3: 3.7 * cm, 4: 0.85 * cm}

        col = cols0[rid] if rid in cols0 else None
        col1 = cols1[rid] if rid in cols1 else None
        cols = tuple([col1, col] + [1.42 * cm] * (col_size - 2))
        # datas = df.reset_index().values.tolist()
        datas = np.array(df.reset_index(drop=True).get_values()).tolist()
        if rid == 5:
            cols = tuple([None] * (col_size - 1) + [3 * cm])
            dfs = pd.pivot_table(df, values='ovccount',
                                 index=['case category'],
                                 columns=['agerange', 'sex'],
                                 aggfunc=sum,
                                 margins=True, margins_name='Total',
                                 fill_value=0)
            # datas = np.vstack((list(dfs), np.array(dfs))).tolist()
            dt_size = len(dfs.index)
            col_size = len(df.columns)
            ds = dt_size + 2
            style = TableStyle(
                [('INNERGRID', (0, 0), (-1, -1), 0.25, colors.black),
                 ('BOX', (0, 0), (-1, -1), 0.5, colors.black),
                 ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
                 ('ALIGN', (1, 3), (-1, -1), 'RIGHT'),
                 ('FONTNAME', (0, ds), (-1, -1), 'Helvetica-Bold'),
                 ('FONTNAME', (col_size - 1, 0), (-1, -1), 'Helvetica-Bold'),
                 ('BACKGROUND', (0, 0), (-1, 0), '#89CFF0')])

            # datas = np.array(dfs.reset_index().get_values()).tolist()
            '''
            pvt_levels = list(dfs.reset_index().columns.levels)
            pvt_labels = list(dfs.reset_index().columns.labels)
            pvt_names = list(dfs.reset_index().columns.names)
            print(pvt_levels, pvt_labels, pvt_names)
            '''
            pvt_names = list(dfs.reset_index().columns.names)
            dcols = {}
            dsize = len(dfs.columns.values)
            pvt_values = list(dfs.reset_index().columns.values)
            # titles = list(pvt_values[0]) + pvt_names
            col_length = 0
            for col in pvt_values:
                col_length = len(col)
                break
            for i in range(0, col_length):
                for col in pvt_values:
                    if i not in dcols:
                        dcols[i] = [col[i]]
                    else:
                        dcols[i].append(col[i])
            fcls = []
            for fc in dcols:
                fcls.append(dcols[fc])
            data_list = dfs.reset_index().values.tolist()
            datas = [[''] + pvt_names] + fcls + data_list
            d0 = 27.86 - (2.22 * dsize)
            cols = tuple([d0 * cm] + [2.22 * cm] * dsize)
        t1 = Table(datas, colWidths=cols)
        t1.setStyle(style)
        element.append(t1)
        # Write the document
        doc = SimpleDocTemplate(
            response, pagesize=A4, rightMargin=20,
            leftMargin=20, topMargin=30, bottomMargin=32,
            keywords="CPIMS, Child Protection in Kenya, UNICEF, DCS, <!NMA!>")
        if rid in [1, 3, 4, 5]:
            doc.pagesize = landscape(A4)
        element.append(Spacer(0.1 * cm, .2 * cm))
        # doc.build(element)
        doc.watermark = 'CPIMS'
        doc.fund_name = ''
        doc.report_info = ''
        doc.source = 'Child Protection Information Management System (CPIMS)'
        doc.build(element, onFirstPage=draw_page, onLaterPages=draw_page,
                  canvasmaker=Canvas)
    except Exception as e:
        raise
    else:
        pass


def get_regions(default_txt=False):
    """Method to return clusters."""
    initial_list = {'': default_txt} if default_txt else {}
    all_list = collections.OrderedDict(initial_list)
    try:
        regions = RegOrgUnit.objects.filter(
            org_unit_type_id='TNRG', is_void=False)
        for a_list in my_list:
            all_list[a_list.id] = a_list.org_unit_name
    except Exception as e:
        error = 'Error getting list - %s' % (str(e))
        print error
        return ()
    else:
        return all_list.items


if __name__ == '__main__':
    pass
